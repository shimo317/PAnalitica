import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import re
import unicodedata
from pathlib import Path

from sklearn.linear_model import LinearRegression
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import train_test_split
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score

# ✅ Word report
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ✅ Toggle global (botón + abrir/cerrar sidebar)
from ui.sidebar_toggle import sidebar_toggle


# ==========================================
# CONFIG
# ==========================================
st.set_page_config(
    page_title="Analítica Predictiva",
    layout="wide",
    initial_sidebar_state="expanded"
)

from auth import init_db, load_authenticated_user, logout_user

init_db()
user = load_authenticated_user()

if not user:
    st.warning("Tu sesión ha expirado o no has iniciado sesión.")
    st.switch_page("Home.py")

if user["role"] not in ["admin", "analista", "usuario"]:
    st.error("No tienes permisos de rol para acceder a esta sección.")
    st.stop()

if "Predictiva" not in user.get("allowed_modules", []):
    st.error("🔒 Módulo bloqueado. No tienes autorización para acceder a 'Predictiva'.")
    st.stop()

# Bootstrap icons
st.markdown(
    '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.1/font/bootstrap-icons.css">',
    unsafe_allow_html=True
)

# ==========================================
# Cargar CSS (global + módulo)
# ==========================================
base_dir = Path(__file__).resolve().parents[1]
global_css = base_dir / "styles" / "global.css"
module_css = base_dir / "styles" / "predictiva.css"

css_text = ""
if global_css.exists():
    css_text += global_css.read_text(encoding="utf-8") + "\n"
if module_css.exists():
    css_text += module_css.read_text(encoding="utf-8") + "\n"

if css_text.strip():
    st.markdown(f"<style>{css_text}</style>", unsafe_allow_html=True)
else:
    st.warning("No se encontró CSS (global.css / predictiva.css).")

# ✅ Toggle después del CSS
sidebar_toggle(user)


# ==========================================
# Helpers
# ==========================================
MESES_ES = ["ene", "feb", "mar", "abr", "may", "jun", "jul", "ago", "sep", "oct", "nov", "dic"]

def money(x):
    try:
        return f"${float(x):,.2f}"
    except Exception:
        return str(x)

def style_dark_matplotlib(ax):
    ax.set_facecolor((0, 0, 0, 0))
    ax.figure.set_facecolor((0, 0, 0, 0))
    ax.tick_params(colors="white")
    ax.xaxis.label.set_color("white")
    ax.yaxis.label.set_color("white")
    ax.title.set_color("white")
    ax.grid(True, alpha=0.18)
    for spine in ax.spines.values():
        spine.set_alpha(0.25)
        spine.set_color("white")

def metrics_pack(y_true, y_pred):
    mae = float(mean_absolute_error(y_true, y_pred))
    rmse = float(np.sqrt(mean_squared_error(y_true, y_pred)))
    r2 = float(r2_score(y_true, y_pred))
    return mae, rmse, r2

def auto_explain(n_meses, mae, rmse, r2, train_r2=None, test_r2=None):
    parts = []
    parts.append(f"Se entrenó el modelo con {n_meses} meses del histórico.")

    if r2 >= 0.80:
        parts.append("El ajuste es alto (R² elevado). La tendencia/patrón se está capturando bien.")
    elif r2 >= 0.50:
        parts.append("El ajuste es medio. Sirve para aproximar, pero hay variación que el modelo no explica.")
    else:
        parts.append("El ajuste es bajo. La serie es variable o corta; la predicción debe tomarse con cautela.")

    parts.append(f"El error promedio estimado (MAE) es {money(mae)}.")
    parts.append(f"El error típico (RMSE) es {money(rmse)}.")

    if train_r2 is not None and test_r2 is not None:
        gap = float(train_r2 - test_r2)
        if gap > 0.25:
            parts.append("Hay señales de sobreajuste: el modelo aprende muy bien el pasado pero generaliza peor.")
        elif gap > 0.10:
            parts.append("Puede haber ligera diferencia entre entrenamiento y prueba; es normal si hay pocos meses.")
        else:
            parts.append("No se observan señales fuertes de sobreajuste en la comparación entrenamiento/prueba.")

    if n_meses < 6:
        parts.append("Recomendación: con menos de 6 meses la proyección puede ser inestable; agrega más histórico si puedes.")
    elif n_meses < 12:
        parts.append("Recomendación: con 6 a 11 meses mejora bastante, pero aún puede haber estacionalidad que no se ve completa.")
    else:
        parts.append("Con 12+ meses es más fácil capturar patrones a lo largo del año.")

    return " ".join(parts)

# ============== WORD HELPERS ==============
def fig_to_png_bytes(fig) -> BytesIO:
    img = BytesIO()
    fig.savefig(img, format="png", dpi=200, bbox_inches="tight")
    img.seek(0)
    return img

def add_df_table_to_doc_full(doc: Document, df: pd.DataFrame, title: str):
    doc.add_heading(title, level=2)
    if df.empty:
        doc.add_paragraph("Sin datos para mostrar.")
        return

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            cells[j].text = str(row[col])

def build_predictive_word_report(
    # Datos y métricas
    total_rows, valid_rows, neg_detected, remove_negative, neg_removed,
    n_meses, mes_ini, mes_fin,
    last_month_str, last_value, fecha_pred_str, pred_next_lin, pred_next_rf,
    best_model_label, best_detail_text,
    met_show: pd.DataFrame,
    interpret_text: str,
    tabla_show: pd.DataFrame,
    # Figuras
    fig_main=None, fig_scatter=None, fig_resid=None
) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Reporte de Analítica Predictiva", level=0)
    doc.add_paragraph("Proyecto: Portal de Analítica (ANALITICA)")
    doc.add_paragraph(f"Rango de meses: {mes_ini} a {mes_fin}")
    doc.add_paragraph("")

    # Estado del archivo
    doc.add_heading("Estado del archivo", level=1)
    doc.add_paragraph(f"Registros cargados: {total_rows}", style="List Bullet")
    doc.add_paragraph(f"Registros válidos (fecha+monto): {valid_rows}", style="List Bullet")
    doc.add_paragraph(f"Montos negativos detectados: {neg_detected}", style="List Bullet")
    doc.add_paragraph(
        f"Quitar negativos: {'Sí' if remove_negative else 'No'} | Negativos removidos: {neg_removed if remove_negative else 0}",
        style="List Bullet"
    )
    doc.add_paragraph("")

    # KPIs
    doc.add_heading("KPIs", level=1)
    doc.add_paragraph(f"Meses analizados: {n_meses}", style="List Bullet")
    doc.add_paragraph(f"Último mes: {last_month_str}", style="List Bullet")
    doc.add_paragraph(f"Último gasto mensual: {money(last_value)}", style="List Bullet")
    doc.add_paragraph(f"Predicción próximo mes (Lineal): {money(pred_next_lin)}", style="List Bullet")
    doc.add_paragraph(f"Predicción próximo mes (Random Forest): {money(pred_next_rf)}", style="List Bullet")
    doc.add_paragraph("")

    # Modelo recomendado
    doc.add_heading("Modelo recomendado", level=1)
    doc.add_paragraph(best_model_label)
    doc.add_paragraph(best_detail_text)
    doc.add_paragraph("")

    # Métricas
    add_df_table_to_doc_full(doc, met_show, "Calidad del modelo (métricas)")
    doc.add_paragraph("")

    # Interpretación
    doc.add_heading("Interpretación automática", level=1)
    doc.add_paragraph(interpret_text)
    doc.add_paragraph(f"Próximo mes estimado ({fecha_pred_str}):")
    doc.add_paragraph(f"• Modelo lineal: {money(pred_next_lin)}")
    doc.add_paragraph(f"• Random Forest: {money(pred_next_rf)}")
    doc.add_paragraph("")

    # Visualizaciones
    doc.add_heading("Visualizaciones", level=1)

    if fig_main is not None:
        doc.add_paragraph("Histórico vs modelos + predicción")
        doc.add_picture(fig_to_png_bytes(fig_main), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

    if fig_scatter is not None:
        doc.add_paragraph("Predicho vs Real (RF)")
        doc.add_picture(fig_to_png_bytes(fig_scatter), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

    if fig_resid is not None:
        doc.add_paragraph("Residuales por mes (RF)")
        doc.add_picture(fig_to_png_bytes(fig_resid), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

    # Tabla final
    add_df_table_to_doc_full(doc, tabla_show, "Serie mensual usada (mensual + modelos + predicción)")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ==========================================
# HERO
# ==========================================
st.markdown(
    """
    <div class="hero-wrap">
        <div class="hero-title">Analítica <span class="accent">Predictiva</span></div>
        <div class="hero-sub">
            Proyección del gasto mensual a partir de tu histórico.
            Se comparan un modelo lineal (explicable) y uno robusto (Random Forest).
        </div>
    </div>
    <div class="sep"></div>
    """,
    unsafe_allow_html=True
)

# ==========================================
# PANEL REQUISITOS
# ==========================================
st.markdown(
    """
    <div class="panel">
        <div style="display:flex; gap:12px;">
            <div style="width:44px;height:44px;border-radius:14px;
                        display:flex;align-items:center;justify-content:center;
                        border:1px solid rgba(34,211,238,.35);
                        background:rgba(2,6,23,.25);color:#22D3EE;">
                <i class="bi bi-graph-up"></i>
            </div>
            <div>
                <div class="panel-title">Requisitos del archivo</div>
                <small>
                    <b>Columnas requeridas:</b>
                    <code>fecha</code>, <code>monto</code>
                    &nbsp;&nbsp;|&nbsp;&nbsp;
                    <b>Opcionales:</b> <code>categoria</code>
                </small>
            </div>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

# ==========================================
# CARGA CSV
# ==========================================
#file = st.file_uploader("Sube tu archivo Excel", type=["xlsx"])
#if not file:
#    st.info("Sube un archivo Excel para comenzar.")
#    st.stop()

#try:
#    df_raw = pd.read_excel(file)
#except Exception as e:
#    st.error(f"No se pudo leer el archivo Excel. Error: {e}")
#    st.stop()


# ==========================================
# UPLOAD + NORMALIZACIÓN MULTI-FUENTE
# ==========================================

import re
import unicodedata

MESES_ES_MAP = {
    "enero": "01", "febrero": "02", "marzo": "03", "abril": "04",
    "mayo": "05", "junio": "06", "julio": "07", "agosto": "08",
    "septiembre": "09", "setiembre": "09", "octubre": "10",
    "noviembre": "11", "diciembre": "12"
}

def normalize_text(txt):
    txt = str(txt).strip().lower()
    txt = unicodedata.normalize("NFKD", txt).encode("ascii", "ignore").decode("ascii")
    txt = re.sub(r"\s+", " ", txt)
    return txt

def first_existing_col(df, candidates):
    norm_cols = {normalize_text(c): c for c in df.columns}
    for c in candidates:
        key = normalize_text(c)
        if key in norm_cols:
            return norm_cols[key]
    return None

def parse_mixed_date(series):
    s = series.astype(str).str.strip()
    dt = pd.to_datetime(s, errors="coerce", dayfirst=True)

    if dt.notna().sum() >= max(3, int(len(s) * 0.5)):
        return dt

    txt = s.str.lower()
    txt = txt.str.replace(r"\s*hs\.?$", "", regex=True)

    for mes, num in MESES_ES_MAP.items():
        txt = txt.str.replace(fr"\b{mes}\b", num, regex=True)

    txt = txt.str.replace(r"\bde\b", " ", regex=True)
    txt = txt.str.replace(r"\s+", " ", regex=True).str.strip()

    return pd.to_datetime(txt, errors="coerce", dayfirst=True)

def detect_source(df):
    cols = {normalize_text(c) for c in df.columns}

    if "# de venta" in cols or "fecha de venta" in cols:
        return "mercadolibre"

    if (
        "paid at" in cols
        or "financial status" in cols
        or "lineitem name" in cols
    ):
        return "shopify"

    if (
        "total (mxn)" in cols
        or "tipo de transaccion" in cols
        or "estatus de la transaccion" in cols
    ):
        return "amazon"

    return "unknown"

def standardize_sales_df(df):
    source = detect_source(df)

    rules = {
        "mercadolibre": {
            "date": ["Fecha de venta"],
            "amount": ["Total (MXN)", "Total MXN", "Total"],
            "product": ["Título de la publicación"],
            "commission": ["Cargo por venta", "Comisión", "Comisiones"],
            "shipping": ["Cargo de envío", "Envío"],
        },

        "amazon": {
            "date": ["Fecha"],
            "amount": ["Total (MXN)", "Total MXN", "Total"],
            "product": ["Detalles del producto"],
            "commission": ["Tarifas de Amazon", "Comisiones"],
            "shipping": ["Otros", "Envío"],
        },

        "shopify": {
            "date": ["Paid at", "Created at"],
            "amount": ["Total", "Total (MXN)"],
            "product": ["Lineitem name"],
            "commission": ["Taxes", "Tax 1 Value"],
            "shipping": ["Shipping", "Shipping Method"],
        },
    }

    if source == "unknown":
        raise ValueError("No se pudo identificar el formato del archivo.")

    # =========================
    # COLUMNAS OBLIGATORIAS
    # =========================
    date_col = first_existing_col(df, rules[source]["date"])
    amount_col = first_existing_col(df, rules[source]["amount"])

    if not date_col or not amount_col:
        raise ValueError(
            f"No se encontraron columnas requeridas para {source}."
        )

    # =========================
    # COLUMNAS OPCIONALES
    # =========================
    product_col = first_existing_col(df, rules[source].get("product", []))
    commission_col = first_existing_col(df, rules[source].get("commission", []))
    shipping_col = first_existing_col(df, rules[source].get("shipping", []))

    # =========================
    # DATAFRAME ESTANDARIZADO
    # =========================
    out = pd.DataFrame()

    # obligatorias
    out["fecha"] = parse_mixed_date(df[date_col])
    out["monto"] = pd.to_numeric(df[amount_col], errors="coerce")

    # opcionales
    out["producto"] = (
        df[product_col].astype(str).str.strip()
        if product_col else np.nan
    )

    out["total_venta"] = out["monto"]

    out["comisiones"] = (
        pd.to_numeric(df[commission_col], errors="coerce")
        if commission_col else 0
    )

    out["envio"] = (
        pd.to_numeric(df[shipping_col], errors="coerce")
        if shipping_col else 0
    )

    out["ingreso_neto"] = (
        out["total_venta"]
        - out["comisiones"].fillna(0)
        - out["envio"].fillna(0)
    )

    out["fuente"] = source

    # =========================
    # LIMPIEZA
    # =========================
    out = out.dropna(subset=["fecha", "monto"]).copy()

    return out

# ==========================================
# UI UPLOAD
# ==========================================

st.markdown('<div class="panel">', unsafe_allow_html=True)

col_up1, col_up2 = st.columns(2)

with col_up1:
    file_excel = st.file_uploader(
        "Excel (.xlsx)",
        type=["xlsx"],
        key="pred_excel"
    )

with col_up2:
    file_csv = st.file_uploader(
        "CSV (.csv)",
        type=["csv"],
        key="pred_csv"
    )

st.markdown('</div>', unsafe_allow_html=True)

# Validación: solo un archivo
if file_excel and file_csv:
    st.warning("⚠️ Sube solo un archivo (Excel o CSV), no ambos.")
    st.stop()

# Archivo activo
file = file_excel if file_excel else file_csv

if not file:
    st.info("Sube un archivo Excel o CSV para comenzar.")
    st.stop()


# ==========================================
# LECTURA ARCHIVO
# ==========================================

if file.name.lower().endswith(".csv"):

    raw_df = None

    # intenta múltiples codificaciones
    for enc in ["utf-8-sig", "utf-8", "latin1", "cp1252"]:
        try:
            file.seek(0)

            # auto detectar separador
            raw_df = pd.read_csv(
                file,
                encoding=enc,
                sep=None,
                engine="python"
            )

            break

        except Exception:
            continue

    if raw_df is None:
        st.error("No se pudo leer el CSV.")
        st.stop()

elif file.name.lower().endswith(".xlsx"):

    file.seek(0)

    try:
        raw_df = pd.read_excel(file)
    except Exception:
        st.error("No se pudo leer el Excel.")
        st.stop()

else:
    st.error("Formato no soportado.")
    st.stop()

# ==========================================
# FIX MERCADO LIBRE
# ==========================================

if detect_source(raw_df) == "unknown" and file.name.lower().endswith(".xlsx"):
    try:
        file.seek(0)
        raw_df = pd.read_excel(file, header=5)
    except Exception:
        pass


# ==========================================
# NORMALIZACIÓN
# ==========================================

try:
    df = standardize_sales_df(raw_df)
except Exception as e:
    st.error(f"No se pudo procesar el archivo: {e}")
    st.stop()

# métricas de limpieza
total_rows = len(raw_df)

valid_rows = int(
    df.dropna(subset=["fecha", "monto"]).shape[0]
)

neg_detected = int(
    (df["monto"] < 0).sum(skipna=True)
)

df = df.dropna(subset=["fecha", "monto"]).copy()

# ==========================================
# Sidebar: Configuración + Validación (pro)
# ==========================================
st.sidebar.markdown("### Configuración")
remove_negative = st.sidebar.checkbox("Quitar montos negativos", value=True)

neg_removed = 0
if remove_negative:
    neg_removed = int((df["monto"] < 0).sum())
    df = df[df["monto"] >= 0].copy()

st.sidebar.markdown("### Estado del archivo")
st.sidebar.write(f"Registros cargados: {total_rows}")
st.sidebar.write(f"Registros válidos (fecha+monto): {valid_rows}")
st.sidebar.write(f"Montos negativos detectados: {neg_detected}")
st.sidebar.write(f"Montos negativos removidos: {neg_removed}" if remove_negative else "Montos negativos removidos: 0 (desactivado)")

if df.empty:
    st.warning("No quedaron datos válidos después de limpiar (fecha, monto).")
    st.stop()

# ==========================================
# CONSTRUIR SERIE MENSUAL
# ==========================================
df["mes"] = df["fecha"].dt.to_period("M").dt.to_timestamp()

serie_mensual = (
    df.groupby("mes", as_index=False)["monto"]
    .sum()
    .rename(columns={"monto": "gasto_mensual"})
    .sort_values("mes")
    .reset_index(drop=True)
)

if len(serie_mensual) < 2:
    st.warning("Necesitas al menos 2 meses distintos para hacer una predicción mensual.")
    st.dataframe(serie_mensual, use_container_width=True, height=280)
    st.stop()

# ==========================================
# FILTROS (PRO) - RANGO POR MES
# ==========================================
st.sidebar.markdown("### Filtros")

meses_opc = serie_mensual["mes"].dt.strftime("%Y-%m").tolist()
mes_ini, mes_fin = st.sidebar.select_slider(
    "Rango de meses",
    options=meses_opc,
    value=(meses_opc[0], meses_opc[-1])
)

start_dt = pd.to_datetime(mes_ini + "-01")
end_dt = pd.to_datetime(mes_fin + "-01")

serie_f = serie_mensual[
    (serie_mensual["mes"] >= start_dt) & (serie_mensual["mes"] <= end_dt)
].copy()

if len(serie_f) < 2:
    st.warning("Con ese filtro quedan menos de 2 meses. Ajusta el rango.")
    st.stop()

serie_f = serie_f.sort_values("mes").reset_index(drop=True)
serie_f["t"] = np.arange(len(serie_f))

X = serie_f[["t"]].values
y = serie_f["gasto_mensual"].values.astype(float)

# ==========================================
# Train/Test
# ==========================================
use_split = len(serie_f) >= 6
if use_split:
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.25, shuffle=False
    )
else:
    X_train, y_train = X, y
    X_test, y_test = None, None

# ==========================================
# MODELOS
# ==========================================
lin = LinearRegression()
lin.fit(X_train, y_train)

rf = RandomForestRegressor(
    n_estimators=400,
    random_state=42,
    max_depth=None
)
rf.fit(X_train, y_train)

# Predicciones para gráfico
y_pred_lin = lin.predict(X)
y_pred_rf = rf.predict(X)

# Métricas train/test
train_mae_lin, train_rmse_lin, train_r2_lin = metrics_pack(y_train, lin.predict(X_train))
train_mae_rf,  train_rmse_rf,  train_r2_rf  = metrics_pack(y_train, rf.predict(X_train))

if use_split:
    test_mae_lin, test_rmse_lin, test_r2_lin = metrics_pack(y_test, lin.predict(X_test))
    test_mae_rf,  test_rmse_rf,  test_r2_rf  = metrics_pack(y_test, rf.predict(X_test))
else:
    test_mae_lin = test_rmse_lin = test_r2_lin = None
    test_mae_rf  = test_rmse_rf  = test_r2_rf  = None

# Próximo mes
next_t = np.array([[int(serie_f["t"].max()) + 1]])
pred_next_lin = float(lin.predict(next_t)[0])
pred_next_rf  = float(rf.predict(next_t)[0])

fecha_pred = (serie_f["mes"].max() + pd.DateOffset(months=1)).to_pydatetime()

# ==========================================
# KPIs
# ==========================================
c1, c2, c3, c4 = st.columns(4)
c1.metric("Meses analizados", f"{len(serie_f)}")
c2.metric("Último mes", serie_f["mes"].max().strftime("%Y-%m"))
c3.metric("Último gasto mensual", money(float(serie_f["gasto_mensual"].iloc[-1])))
c4.metric("Predicción próximo mes (RF)", money(pred_next_rf))

st.divider()

# ==========================================
# TABLA MÉTRICAS
# ==========================================
st.markdown("## Calidad del modelo (métricas)")
st.caption(
    "Train = entrenamiento (lo que vio el modelo). "
    "Test = prueba (meses que NO vio; es lo más realista para evaluar)."
)

with st.expander("¿Qué significan MAE, RMSE y R²? (explicación simple)", expanded=False):
    st.write(
        "MAE: promedio del error en pesos. Si MAE = $500, el modelo se equivoca ~ $500 en promedio.\n\n"
        "RMSE: parecido al MAE pero castiga más los errores grandes. Si hay meses con picos fuertes, RMSE sube.\n\n"
        "R²: qué tanto explica el modelo la variación del gasto.\n"
        "• 1.00 = perfecto\n"
        "• 0.50 = explica la mitad\n"
        "• 0.00 = igual que usar el promedio\n"
        "• negativo = peor que usar el promedio (serie muy variable o pocos datos)\n\n"
        "Tip: para comparar modelos, fíjate principalmente en MAE/RMSE en TEST."
    )

met_rows = [
    {
        "Modelo": "Lineal",
        "MAE (train)": train_mae_lin,
        "RMSE (train)": train_rmse_lin,
        "R² (train)": train_r2_lin,
        "MAE (test)": (np.nan if not use_split else test_mae_lin),
        "RMSE (test)": (np.nan if not use_split else test_rmse_lin),
        "R² (test)": (np.nan if not use_split else test_r2_lin),
    },
    {
        "Modelo": "Random Forest",
        "MAE (train)": train_mae_rf,
        "RMSE (train)": train_rmse_rf,
        "R² (train)": train_r2_rf,
        "MAE (test)": (np.nan if not use_split else test_mae_rf),
        "RMSE (test)": (np.nan if not use_split else test_rmse_rf),
        "R² (test)": (np.nan if not use_split else test_r2_rf),
    },
]
met_df = pd.DataFrame(met_rows)

def fmt_money_or_dash(v):
    if pd.isna(v): return "-"
    return money(v)

def fmt_r2_or_dash(v):
    if pd.isna(v): return "-"
    return f"{float(v):.3f}"

met_show = met_df.copy()
for c in ["MAE (train)", "RMSE (train)", "MAE (test)", "RMSE (test)"]:
    met_show[c] = met_show[c].apply(fmt_money_or_dash)
for c in ["R² (train)", "R² (test)"]:
    met_show[c] = met_show[c].apply(fmt_r2_or_dash)

st.dataframe(
    met_show,
    use_container_width=True,
    hide_index=True,
    height=140
)

if not use_split:
    st.info("No hay suficientes meses para separar Train/Test. Agrega más meses para métricas más realistas.")

# Modelo recomendado (pro)
if use_split:
    best = "Random Forest" if test_mae_rf < test_mae_lin else "Lineal"
    best_mae = test_mae_rf if best == "Random Forest" else test_mae_lin
    best_model_label = f"Modelo recomendado: {best}"
    best_detail = f"Se eligió por menor MAE test = {money(best_mae)}."
    st.success(f"{best_model_label} ({best_detail})")
else:
    best = "Random Forest" if train_mae_rf < train_mae_lin else "Lineal"
    best_mae = train_mae_rf if best == "Random Forest" else train_mae_lin
    best_model_label = f"Modelo sugerido: {best}"
    best_detail = f"Se eligió por menor MAE train = {money(best_mae)}. Recomendado: usar 6+ meses para Train/Test."
    st.warning(f"{best_model_label} ({best_detail})")

# ==========================================
# GRÁFICAS + TEXTO
# ==========================================
colA, colB = st.columns(2)

with colA:
    st.markdown("## Histórico vs modelos + predicción")
    fig, ax = plt.subplots(figsize=(7.6, 4.4))
    ax.plot(serie_f["mes"], y, marker="o", linewidth=2, label="Histórico")
    ax.plot(serie_f["mes"], y_pred_lin, linestyle="--", linewidth=2, label="Ajuste lineal")
    ax.plot(serie_f["mes"], y_pred_rf, linestyle="--", linewidth=2, label="Ajuste RF")
    ax.scatter(fecha_pred, pred_next_lin, s=60, label="Pred. lineal")
    ax.scatter(fecha_pred, pred_next_rf,  s=60, label="Pred. RF")

    ax.set_xlabel("Mes")
    ax.set_ylabel("Gasto mensual")
    ax.set_title("Comparación de modelos y proyección")
    plt.xticks(rotation=35, ha="right")

    style_dark_matplotlib(ax)
    ax.legend()
    plt.tight_layout()
    st.pyplot(fig, use_container_width=True)

with colB:
    r2_main = train_r2_rf if not use_split else test_r2_rf
    mae_main = train_mae_rf if not use_split else test_mae_rf
    rmse_main = train_rmse_rf if not use_split else test_rmse_rf

    txt = auto_explain(
        n_meses=len(serie_f),
        mae=mae_main,
        rmse=rmse_main,
        r2=r2_main,
        train_r2=train_r2_rf,
        test_r2=None if not use_split else test_r2_rf
    )

    st.markdown(
        f"""
        <div class="interpret-card">
            <div class="interpret-title">Interpretación automática</div>
            <div class="interpret-text">
                {txt}<br><br>
                <b>Próximo mes estimado ({fecha_pred.strftime('%Y-%m')}):</b><br>
                Modelo lineal: {money(pred_next_lin)}<br>
                Random Forest: {money(pred_next_rf)}
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

st.divider()

# ==========================================
# DIAGNÓSTICO RF
# ==========================================
st.markdown("## Diagnóstico del modelo (Random Forest)")

colC, colD = st.columns(2)

with colC:
    fig2, ax2 = plt.subplots(figsize=(7.6, 4.2))
    ax2.scatter(y, y_pred_rf)
    ax2.set_xlabel("Real")
    ax2.set_ylabel("Predicho (RF)")
    ax2.set_title("Predicho vs Real (RF)")
    style_dark_matplotlib(ax2)
    plt.tight_layout()
    st.pyplot(fig2, use_container_width=True)

with colD:
    residuals = y - y_pred_rf
    fig3, ax3 = plt.subplots(figsize=(7.6, 4.2))
    ax3.axhline(0, linewidth=2, alpha=0.6)
    ax3.plot(serie_f["mes"], residuals, marker="o")
    ax3.set_xlabel("Mes")
    ax3.set_ylabel("Error (Real - Predicho)")
    ax3.set_title("Residuales por mes (RF)")
    plt.xticks(rotation=35, ha="right")
    style_dark_matplotlib(ax3)
    plt.tight_layout()
    st.pyplot(fig3, use_container_width=True)

st.divider()

# ==========================================
# TABLA + DESCARGAS
# ==========================================
st.markdown('<div class="section-title">Serie mensual usada</div>', unsafe_allow_html=True)

tabla = serie_f[["mes", "gasto_mensual"]].copy()
tabla["mes"] = tabla["mes"].dt.strftime("%Y-%m")
tabla["pred_lineal"] = y_pred_lin
tabla["pred_rf"] = y_pred_rf
tabla["residual_rf"] = (y - y_pred_rf)

tabla_show = tabla.copy()
tabla_show["gasto_mensual"] = tabla_show["gasto_mensual"].apply(money)
tabla_show["pred_lineal"] = tabla_show["pred_lineal"].apply(money)
tabla_show["pred_rf"] = tabla_show["pred_rf"].apply(money)
tabla_show["residual_rf"] = tabla_show["residual_rf"].apply(money)

st.dataframe(tabla_show, use_container_width=True, height=320)

# Export con fila próximo mes
export_final = tabla.copy()
export_final["gasto_mensual"] = export_final["gasto_mensual"].round(4)
export_final["pred_lineal"] = export_final["pred_lineal"].round(4)
export_final["pred_rf"] = export_final["pred_rf"].round(4)
export_final["residual_rf"] = export_final["residual_rf"].round(4)

pred_row = pd.DataFrame([{
    "mes": fecha_pred.strftime("%Y-%m"),
    "gasto_mensual": "",
    "pred_lineal": round(pred_next_lin, 4),
    "pred_rf": round(pred_next_rf, 4),
    "residual_rf": ""
}])

# export_final = pd.concat([export_final, pred_row], ignore_index=True)
# csv_bytes = export_final.to_csv(index=False).encode("utf-8")

# st.download_button(
#    "⬇ Descargar CSV (mensual + modelos + predicción)",
 #   data=csv_bytes,
  #  file_name="reporte_predictivo_mensual_pro.csv",
   # mime="text/csv"
#)

# ==========================================
# WORD (GENERAR + DESCARGAR)
# ==========================================
st.markdown(
    """
    <div class="word-card">
      <div class="word-title">Descargar Reporte Word</div>
      <div class="word-sub">
        Genera un archivo .docx con KPIs, métricas, interpretación, predicción y gráficas.
      </div>
    </div>
    """,
    unsafe_allow_html=True
)

if "word_pred_bytes" not in st.session_state:
    st.session_state.word_pred_bytes = None

colW1, colW2 = st.columns([1, 1.2])

with colW1:
    generar_word = st.button("📄 Generar Word (predictiva)")

with colW2:
    if st.session_state.word_pred_bytes:
        st.download_button(
            "⬇ Descargar Reporte Word (.docx)",
            data=st.session_state.word_pred_bytes,
            file_name="reporte_predictivo_completo.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.markdown("<div style='height: 52px;'></div>", unsafe_allow_html=True)

if generar_word:
    st.session_state.word_pred_bytes = build_predictive_word_report(
        total_rows=total_rows,
        valid_rows=valid_rows,
        neg_detected=neg_detected,
        remove_negative=remove_negative,
        neg_removed=neg_removed,
        n_meses=len(serie_f),
        mes_ini=mes_ini,
        mes_fin=mes_fin,
        last_month_str=serie_f["mes"].max().strftime("%Y-%m"),
        last_value=float(serie_f["gasto_mensual"].iloc[-1]),
        fecha_pred_str=fecha_pred.strftime("%Y-%m"),
        pred_next_lin=pred_next_lin,
        pred_next_rf=pred_next_rf,
        best_model_label=best_model_label,
        best_detail_text=best_detail,
        met_show=met_show,
        interpret_text=txt,
        tabla_show=tabla_show,
        fig_main=fig,
        fig_scatter=fig2,
        fig_resid=fig3
    )

st.caption("© 2025 Portal de Analítica | Módulo Predictivo")
