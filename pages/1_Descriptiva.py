import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import re
import unicodedata 
from pathlib import Path
from datetime import timedelta


from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ui.sidebar_toggle import sidebar_toggle
from auth import init_db, load_authenticated_user, logout_user

# ==========================================
# CONFIG
# ==========================================
st.set_page_config(
    page_title="Analítica Descriptiva",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# AUTH
# ==========================================
init_db()
user = load_authenticated_user()

if not user:
    st.warning("Tu sesión ha expirado o no has iniciado sesión.")
    st.switch_page("Home.py")

if user["role"] not in ["admin", "analista", "usuario"]:
    st.error("No tienes permisos para acceder a esta sección.")
    st.stop()

if "Descriptiva" not in user.get("allowed_modules", []):
    st.error("🔒 Módulo bloqueado. No tienes autorización para acceder a 'Descriptiva'.")
    st.stop()

# Bootstrap icons
st.markdown(
    '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.1/font/bootstrap-icons.css">',
    unsafe_allow_html=True
)

# ==========================================
# Cargar CSS (global + módulo)
# ==========================================
base_dir = Path(__file__).resolve().parents[1]
global_css = base_dir / "styles" / "global.css"
module_css = base_dir / "styles" / "descriptiva.css"

css_text = ""
if global_css.exists():
    css_text += global_css.read_text(encoding="utf-8") + "\n"
if module_css.exists():
    css_text += module_css.read_text(encoding="utf-8") + "\n"

css_text += """
.stCaption, .stCaption *{ color: rgba(226,232,240,0.88) !important; }
.stMarkdown, .stMarkdown *{ color: rgba(255,255,255,0.95); }

.auth-topbar{
    display:flex;
    justify-content:space-between;
    align-items:center;
    gap:14px;
    margin: 4px 0 12px 0;
    padding: 10px 14px;
    border: 1px solid rgba(255,255,255,.08);
    border-radius: 16px;
    background: rgba(255,255,255,.03);
    backdrop-filter: blur(6px);
}
.auth-user{
    color: rgba(255,255,255,.92);
    font-size: 0.95rem;
}
"""

if css_text.strip():
    st.markdown(f"<style>{css_text}</style>", unsafe_allow_html=True)
else:
    st.warning("No se encontró CSS (global.css / descriptiva.css).")

sidebar_toggle(user)

# ==========================================
# HEADER USUARIO / SESIÓN
# ==========================================
c1, c2 = st.columns([0.82, 0.18])
with c1:
    st.markdown(
        f"""
        <div class="auth-topbar">
            <div class="auth-user">
                👤 <b>{user['username']}</b> &nbsp;|&nbsp; Rol: <b>{user['role'].replace('analista', 'Operación').replace('usuario', 'Ventas').replace('admin', 'Administrador')}</b>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )
with c2:
    if st.button("Cerrar sesión", use_container_width=True):
        logout_user()
        st.switch_page("Home.py")

MESES_ES = ["ene", "feb", "mar", "abr", "may", "jun", "jul", "ago", "sep", "oct", "nov", "dic"]

def money(x) -> str:
    try:
        return f"${float(x):,.2f}"
    except Exception:
        return str(x)

def style_dark_matplotlib(ax):
    ax.set_facecolor((0, 0, 0, 0))
    ax.figure.set_facecolor((0, 0, 0, 0))
    ax.tick_params(colors="white")
    ax.xaxis.label.set_color("white")
    ax.yaxis.label.set_color("white")
    ax.title.set_color("white")
    ax.grid(True, alpha=0.18)
    for spine in ax.spines.values():
        spine.set_alpha(0.25)
        spine.set_color("white")

def format_date_es(d):
    if pd.isna(d):
        return ""
    d = pd.to_datetime(d).date()
    return f"{d.day:02d} {MESES_ES[d.month-1]} {d.year}"

def make_period_labels(df: pd.DataFrame, nivel: str) -> pd.DataFrame:
    out = df.copy()
    if nivel == "Mensual":
        out["periodo"] = out["fecha"].dt.to_period("M").astype(str)
        y = out["fecha"].dt.year
        m = out["fecha"].dt.month
        out["periodo_etq"] = m.map(lambda mm: MESES_ES[mm-1]) + " " + y.astype(str)
        return out

    iso = out["fecha"].dt.isocalendar()
    out["iso_year"] = iso["year"].astype(int)
    out["iso_week"] = iso["week"].astype(int)

    start = out["fecha"] - pd.to_timedelta(out["fecha"].dt.weekday, unit="D")
    end = start + pd.to_timedelta(6, unit="D")

    out["semana_inicio"] = start.dt.date
    out["semana_fin"] = end.dt.date
    out["periodo"] = out["iso_year"].astype(str) + "-W" + out["iso_week"].astype(str).str.zfill(2)

    out["periodo_etq"] = out.apply(
        lambda r: f"Semana {int(r['iso_week'])} ({format_date_es(r['semana_inicio'])[:-5]} – {format_date_es(r['semana_fin'])[:-5]})",
        axis=1
    )
    return out

def build_insights(resumen_periodo: pd.DataFrame, resumen_categoria: pd.Series, unidad: str) -> dict:
    insights = {}

    m = resumen_periodo.copy()
    m["gasto_total"] = pd.to_numeric(m["gasto_total"], errors="coerce").fillna(0)

    n = int(len(m))
    unidad_txt = "mes" if unidad == "mes" else "semana"
    unidad_pl = "meses" if unidad == "mes" else "semanas"

    first = float(m["gasto_total"].iloc[0])
    last = float(m["gasto_total"].iloc[-1])
    cambio_pct = ((last - first) / first * 100) if first != 0 else 0.0

    x = np.arange(n)
    y = m["gasto_total"].values.astype(float)
    slope_pct = 0.0
    if n >= 2:
        slope = np.polyfit(x, y, 1)[0]
        y_mean = float(np.mean(y)) if float(np.mean(y)) != 0 else 1.0
        slope_pct = (float(slope) / y_mean) * 100

    if slope_pct > 2:
        trend_txt = "tendencia al alza"
    elif slope_pct < -2:
        trend_txt = "tendencia a la baja"
    else:
        trend_txt = "tendencia estable"

    vol = (float(np.std(y)) / (float(np.mean(y)) if float(np.mean(y)) != 0 else 1.0)) * 100
    if vol >= 25:
        vol_txt = "alta variabilidad (gasto irregular)"
    elif vol >= 12:
        vol_txt = "variabilidad moderada"
    else:
        vol_txt = "gasto relativamente constante"

    i_max = int(m["gasto_total"].values.argmax())
    i_min = int(m["gasto_total"].values.argmin())

    periodo_max = str(m["periodo_etq"].iloc[i_max])
    val_max = float(m["gasto_total"].iloc[i_max])

    periodo_min = str(m["periodo_etq"].iloc[i_min])
    val_min = float(m["gasto_total"].iloc[i_min])

    jump_txt = ""
    if n >= 3:
        dif = pd.Series(y).pct_change().replace([np.inf, -np.inf], np.nan).dropna()
        if not dif.empty:
            j = int(dif.abs().values.argmax())
            idx_to = int(dif.index[j])
            idx_from = idx_to - 1
            salto = float(dif.iloc[j] * 100)
            if abs(salto) >= 20:
                p_from = str(m["periodo_etq"].iloc[idx_from])
                p_to = str(m["periodo_etq"].iloc[idx_to])
                jump_txt = f"Se detecta un cambio brusco entre {p_from} → {p_to} ({salto:+.1f}%)."

    insights["periodo"] = (
        f"En el periodo analizado ({n} {unidad_pl}), el gasto presenta {trend_txt}. "
        f"Del primer al último {unidad_txt} hay un cambio de {cambio_pct:+.1f}%. "
        f"Además, se observa {vol_txt}. "
        f"El mayor gasto fue en {periodo_max} ({money(val_max)}), y el menor en {periodo_min} ({money(val_min)}). "
        f"{jump_txt}".strip()
    )

    cat = resumen_categoria.copy()
    total = float(cat.sum()) if float(cat.sum()) != 0 else 1.0

    top1 = str(cat.index[0]) if len(cat) else "N/A"
    top1_val = float(cat.iloc[0]) if len(cat) else 0.0
    top1_pct = (top1_val / total) * 100

    top3_val = float(cat.head(3).sum()) if len(cat) >= 3 else float(cat.sum())
    top3_pct = (top3_val / total) * 100

    if top1_pct >= 40:
        conc_txt = "El gasto está muy concentrado en una sola categoría."
    elif top1_pct >= 25:
        conc_txt = "El gasto está concentrado en la categoría principal."
    else:
        conc_txt = "El gasto está distribuido entre varias categorías."

    insights["categoria"] = (
        f"La categoría con mayor gasto es {top1} con {money(top1_val)} ({top1_pct:.1f}% del total). "
        f"Las 3 principales categorías representan {top3_pct:.1f}% del gasto. "
        f"{conc_txt}"
    )

    return insights

def format_table_for_display(df: pd.DataFrame, money_cols=None) -> pd.DataFrame:
    money_cols = money_cols or []
    out = df.copy()

    for c in money_cols:
        if c in out.columns:
            out[c] = out[c].apply(money)

    return out

def fig_to_png_bytes(fig) -> BytesIO:
    img = BytesIO()
    fig.savefig(img, format="png", dpi=200, bbox_inches="tight")
    img.seek(0)
    return img

def add_df_table_to_doc_full(doc: Document, df: pd.DataFrame, title: str, money_cols=None):
    money_cols = set(money_cols or [])
    doc.add_heading(title, level=2)

    if df.empty:
        doc.add_paragraph("Sin datos para mostrar.")
        return

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            if col in money_cols:
                cells[j].text = money(val)
            else:
                cells[j].text = str(val)

def build_word_report_full(
    df_filtrado: pd.DataFrame,
    resumen_periodo_show: pd.DataFrame,
    resumen_categoria: pd.Series,
    insights: dict,
    nivel: str,
    start_date,
    end_date,
    fig_evolucion=None,
    fig_categorias=None
) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Reporte de Analítica Descriptiva", level=0)
    p = doc.add_paragraph("Proyecto: Portal de Analítica (ANALITICA)\n")
    p.add_run(f"Nivel de análisis: {nivel}\n")
    p.add_run(f"Rango de fechas: {start_date} a {end_date}\n")
    doc.add_paragraph("")

    doc.add_heading("KPIs", level=1)
    gasto_total = float(df_filtrado["monto"].sum())
    promedio = float(df_filtrado["monto"].mean())
    maximo = float(df_filtrado["monto"].max())
    registros = int(len(df_filtrado))

    for line in [
        f"Gasto total: {money(gasto_total)}",
        f"Promedio: {money(promedio)}",
        f"Máximo: {money(maximo)}",
        f"Registros: {registros}",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph("")
    doc.add_heading("Interpretación automática", level=1)
    doc.add_paragraph(insights.get("periodo", ""))
    doc.add_paragraph(insights.get("categoria", ""))
    doc.add_paragraph("")

    doc.add_heading("Visualizaciones", level=1)

    if fig_evolucion is not None:
        doc.add_paragraph("Evolución del gasto")
        img1 = fig_to_png_bytes(fig_evolucion)
        doc.add_picture(img1, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

    if fig_categorias is not None:
        doc.add_paragraph("Gasto total por categoría")
        img2 = fig_to_png_bytes(fig_categorias)
        doc.add_picture(img2, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")

    add_df_table_to_doc_full(
        doc,
        resumen_periodo_show,
        title="Resumen del periodo",
        money_cols=["gasto_total", "gasto_promedio"]
    )

    doc.add_heading("Ranking por categoría", level=1)
    cat_df = resumen_categoria.reset_index()
    cat_df.columns = ["categoria", "monto_total"]
    add_df_table_to_doc_full(
        doc,
        cat_df,
        title="Ranking completo por categoría",
        money_cols=["monto_total"]
    )

    add_df_table_to_doc_full(
        doc,
        df_filtrado,
        title="Detalle completo de movimientos filtrados",
        money_cols=["monto"]
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

st.markdown(
    """
    <div class="hero-wrap">
        <div class="hero-title">Analítica <span class="accent">Descriptiva</span></div>
        <div class="hero-sub">
            Sube tu archivo Excel o CSV, filtra por fechas y categorías, visualiza KPIs y revisa insights automáticos.
        </div>
    </div>
    <div class="sep"></div>
    """,
    unsafe_allow_html=True
)

st.markdown(
    """
    <div class="panel">
        <div class="panel-row">
            <div class="panel-icon">
                <i class="bi bi-info-circle"></i>
            </div>
            <div>
                <div class="panel-title">Requisitos del archivo</div>
                <small>
                    <b>Columnas requeridas:</b>
                    <code>fecha</code>, <code>categoria</code>, <code>monto</code>
                    &nbsp;&nbsp;|&nbsp;&nbsp;
                    <b>Opcionales:</b>
                    <code>metodo_pago</code>, <code>nota</code>
                </small>
            </div>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

#file = st.file_uploader("Sube tu archivo Excel", type=["xlsx"])
#if not file:
 #   st.info("Sube un archivo Excel para comenzar.")
 #  st.stop()



MESES_ES_MAP = {
    "enero": "01", "febrero": "02", "marzo": "03", "abril": "04",
    "mayo": "05", "junio": "06", "julio": "07", "agosto": "08",
    "septiembre": "09", "setiembre": "09", "octubre": "10",
    "noviembre": "11", "diciembre": "12"
}

def normalize_text(txt):
    txt = str(txt).strip().lower()
    txt = unicodedata.normalize("NFKD", txt).encode("ascii", "ignore").decode("ascii")
    txt = re.sub(r"\s+", " ", txt)
    return txt

def first_existing_col(df, candidates):
    norm_cols = {normalize_text(c): c for c in df.columns}
    for c in candidates:
        key = normalize_text(c)
        if key in norm_cols:
            return norm_cols[key]
    return None

def parse_mixed_date(series):
    s = series.astype(str).str.strip()

    # intento 1: formato normal
    dt = pd.to_datetime(s, errors="coerce", dayfirst=True)

    # si funciona suficiente, lo regresamos
    if dt.notna().sum() >= max(3, int(len(s) * 0.5)):
        return dt

    # intento 2: fechas españolas tipo "30 de abril de 2026 21:46 hs."
    txt = s.str.lower()
    txt = txt.str.replace(r"\s*hs\.?$", "", regex=True)

    for mes, num in MESES_ES_MAP.items():
        txt = txt.str.replace(fr"\b{mes}\b", num, regex=True)

    txt = txt.str.replace(r"\bde\b", " ", regex=True)
    txt = txt.str.replace(r"\s+", " ", regex=True).str.strip()

    return pd.to_datetime(txt, errors="coerce", dayfirst=True)

def detect_source(df):
    cols = {normalize_text(c) for c in df.columns}

    if "# de venta" in cols or "fecha de venta" in cols:
        return "mercadolibre"
    if "paid at" in cols or "financial status" in cols or "lineitem name" in cols:
        return "shopify"
    if "total (mxn)" in cols or "tipo de transaccion" in cols:
        return "amazon"

    return "unknown"

def standardize_sales_df(df):
    source = detect_source(df)

    rules = {
        "mercadolibre": {
            "date": ["Fecha de venta"],
            "amount": ["Total (MXN)"],
            "category": ["Título de la publicación", "Canal de venta", "Estado"],

            # OPCIONALES
            "product": ["Título de la publicación"],
            "commissions": ["Cargo por venta y envíos", "Comisión de Mercado Libre"],
            "shipping": ["Costo de envío"],
        },

        "amazon": {
            "date": ["Fecha"],
            "amount": ["Total (MXN)"],
            "category": ["Detalles del producto", "Tipo de transacción", "Estatus de la transacción"],

            # OPCIONALES
            "product": ["Detalles del producto"],
            "commissions": ["Tarifas de Amazon"],
            "shipping": ["Otros"],
        },

        "shopify": {
            "date": ["Paid at", "Created at"],
            "amount": ["Total"],
            "category": ["Lineitem name", "Vendor", "Shipping Method"],

            # OPCIONALES
            "product": ["Lineitem name"],
            "commissions": ["Fees", "Processing Fees"],
            "shipping": ["Shipping"],
        },
    }

    if source == "unknown":
        raise ValueError("No se pudo identificar el formato del archivo.")

    # =====================================
    # COLUMNAS OBLIGATORIAS
    # =====================================

    date_col = first_existing_col(df, rules[source]["date"])
    amount_col = first_existing_col(df, rules[source]["amount"])
    category_col = first_existing_col(df, rules[source]["category"])

    if not date_col or not amount_col or not category_col:
        raise ValueError(
            f"No se encontraron columnas requeridas para el formato {source}."
        )

    # =====================================
    # COLUMNAS OPCIONALES
    # =====================================

    product_col = first_existing_col(
        df,
        rules[source].get("product", [])
    )

    commissions_col = first_existing_col(
        df,
        rules[source].get("commissions", [])
    )

    shipping_col = first_existing_col(
        df,
        rules[source].get("shipping", [])
    )

    # =====================================
    # DATAFRAME FINAL
    # =====================================

    out = pd.DataFrame()

    # OBLIGATORIAS
    out["fecha"] = parse_mixed_date(df[date_col])

    out["monto"] = pd.to_numeric(
        df[amount_col],
        errors="coerce"
    )

    out["categoria"] = (
        df[category_col]
        .astype(str)
        .str.strip()
    )

    # =====================================
    # OPCIONALES
    # =====================================

    out["producto"] = (
        df[product_col].astype(str).str.strip()
        if product_col
        else "Sin producto"
    )

    out["venta_total"] = pd.to_numeric(
        df[amount_col],
        errors="coerce"
    ).fillna(0)

    out["comisiones"] = (
        pd.to_numeric(df[commissions_col], errors="coerce").fillna(0)
        if commissions_col
        else 0
    )

    out["envio"] = (
        pd.to_numeric(df[shipping_col], errors="coerce").fillna(0)
        if shipping_col
        else 0
    )

    out["ingreso_neto"] = (
        out["venta_total"]
        - out["comisiones"]
        - out["envio"]
    )

    # TU SISTEMA ACTUAL SIGUE USANDO "monto"
    out["monto"] = out["ingreso_neto"]

    # plataforma origen
    out["plataforma"] = source

    # =====================================
    # LIMPIEZA
    # =====================================

    out["categoria"] = out["categoria"].replace(
        {
            "nan": np.nan,
            "None": np.nan,
            "": np.nan,
            "NA": np.nan,
            "N/A": np.nan
        }
    )

    out = out.dropna(
        subset=["fecha", "categoria", "monto"]
    ).copy()

    return out

# ==========================================
# UPLOAD + NORMALIZACIÓN MULTI-FUENTE
# ==========================================

import re
import unicodedata

MESES_ES_MAP = {
    "enero": "01", "febrero": "02", "marzo": "03", "abril": "04",
    "mayo": "05", "junio": "06", "julio": "07", "agosto": "08",
    "septiembre": "09", "setiembre": "09", "octubre": "10",
    "noviembre": "11", "diciembre": "12"
}

def normalize_text(txt):
    txt = str(txt).strip().lower()
    txt = unicodedata.normalize("NFKD", txt).encode("ascii", "ignore").decode("ascii")
    txt = re.sub(r"\s+", " ", txt)
    return txt

def first_existing_col(df, candidates):
    norm_cols = {normalize_text(c): c for c in df.columns}
    for c in candidates:
        key = normalize_text(c)
        if key in norm_cols:
            return norm_cols[key]
    return None

def parse_mixed_date(series):
    s = series.astype(str).str.strip()

    dt = pd.to_datetime(s, errors="coerce", dayfirst=True)

    if dt.notna().sum() >= max(3, int(len(s) * 0.5)):
        return dt

    txt = s.str.lower()
    txt = txt.str.replace(r"\s*hs\.?$", "", regex=True)

    for mes, num in MESES_ES_MAP.items():
        txt = txt.str.replace(fr"\b{mes}\b", num, regex=True)

    txt = txt.str.replace(r"\bde\b", " ", regex=True)
    txt = txt.str.replace(r"\s+", " ", regex=True).str.strip()

    return pd.to_datetime(txt, errors="coerce", dayfirst=True)

def detect_source(df):
    cols = {normalize_text(c) for c in df.columns}

    if "# de venta" in cols or "fecha de venta" in cols:
        return "mercadolibre"
    if "paid at" in cols or "financial status" in cols or "lineitem name" in cols:
        return "shopify"
    if "total (mxn)" in cols or "tipo de transaccion" in cols:
        return "amazon"

    return "unknown"

def standardize_sales_df(df):
    source = detect_source(df)

    rules = {
        "mercadolibre": {
            "date": ["Fecha de venta"],
            "amount": ["Total (MXN)"],
            "category": ["Título de la publicación", "Canal de venta", "Estado"],
        },
        "amazon": {
            "date": ["Fecha"],
            "amount": ["Total (MXN)"],
            "category": ["Detalles del producto", "Tipo de transacción", "Estatus de la transacción"],
        },
        "shopify": {
            "date": ["Paid at", "Created at"],
            "amount": ["Total"],
            "category": ["Lineitem name", "Vendor", "Shipping Method"],
        },
    }

    if source == "unknown":
        raise ValueError("No se pudo identificar el formato del archivo.")

    date_col = first_existing_col(df, rules[source]["date"])
    amount_col = first_existing_col(df, rules[source]["amount"])
    category_col = first_existing_col(df, rules[source]["category"])

    if not date_col or not amount_col or not category_col:
        raise ValueError(f"No se encontraron columnas requeridas para el formato {source}.")

    out = pd.DataFrame()
    out["fecha"] = parse_mixed_date(df[date_col])
    out["monto"] = pd.to_numeric(df[amount_col], errors="coerce")
    out["categoria"] = df[category_col].astype(str).str.strip()
    out["fuente"] = source

    out["categoria"] = out["categoria"].replace(
        {"nan": np.nan, "None": np.nan, "": np.nan, "NA": np.nan, "N/A": np.nan}
    )

    out = out.dropna(subset=["fecha", "categoria", "monto"]).copy()
    return out


# ==========================================
# UI UPLOAD
# ==========================================

st.markdown('<div class="panel">', unsafe_allow_html=True)

uploaded_file = st.file_uploader(
    "Sube tu archivo de ventas (Excel o CSV)",
    type=["xlsx", "csv"]
)

st.markdown('</div>', unsafe_allow_html=True)

if not uploaded_file:
    st.info("Sube un archivo Excel o CSV para comenzar.")
    st.stop()


# ==========================================
# LECTURA ARCHIVO
# ==========================================

if uploaded_file.name.lower().endswith(".csv"):
    raw_df = None
    for enc in ["utf-8-sig", "utf-8", "latin1", "cp1252"]:
        try:
            uploaded_file.seek(0)
            raw_df = pd.read_csv(uploaded_file, encoding=enc)
            break
        except Exception:
            continue

    if raw_df is None:
        st.error("No se pudo leer el CSV.")
        st.stop()

elif uploaded_file.name.lower().endswith(".xlsx"):
    uploaded_file.seek(0)

    try:
        raw_df = pd.read_excel(uploaded_file)
    except Exception:
        st.error("No se pudo leer el archivo Excel.")
        st.stop()

else:
    st.error("Formato no soportado.")
    st.stop()


# ==========================================
# FIX MERCADO LIBRE (HEADER DESPLAZADO)
# ==========================================

if detect_source(raw_df) == "unknown" and uploaded_file.name.lower().endswith(".xlsx"):
    try:
        uploaded_file.seek(0)
        raw_df = pd.read_excel(uploaded_file, sheet_name=0, header=5)
    except Exception:
        pass


# ==========================================
# NORMALIZACIÓN FINAL
# ==========================================

try:
    df = standardize_sales_df(raw_df)
except Exception as e:
    st.error(f"No se pudo estandarizar el archivo: {e}")
    st.stop()

if df.empty:
    st.warning("El archivo no tiene datos válidos después de estandarizar.")
    st.stop()

required = {"fecha", "categoria", "monto"}
if not required.issubset(set(df.columns)):
    st.error(f"El archivo Excel o CSV debe contener al menos estas columnas: {', '.join(sorted(required))}.")
    st.stop()

df["fecha"] = pd.to_datetime(df["fecha"], errors="coerce")
df["monto"] = pd.to_numeric(df["monto"], errors="coerce")
df["categoria"] = df["categoria"].astype(str).str.strip()
df = df.dropna(subset=["fecha", "categoria", "monto"])

if df.empty:
    st.warning("El archivo no tiene datos válidos después de limpiar (fecha/categoria/monto).")
    st.stop()

st.sidebar.markdown("### Filtros")

rango = st.sidebar.date_input(
    "Rango de fechas",
    (df["fecha"].min().date(), df["fecha"].max().date())
)
start, end = rango

df_f = df[(df["fecha"].dt.date >= start) & (df["fecha"].dt.date <= end)].copy()

cats = sorted(df_f["categoria"].dropna().unique())
sel_cats = st.sidebar.multiselect("Categorías", cats, default=cats)
df_f = df_f[df_f["categoria"].isin(sel_cats)].copy()

if df_f.empty:
    st.warning("Con los filtros actuales no hay datos para mostrar. Ajusta el rango o categorías.")
    st.stop()

c1, c2, c3, c4 = st.columns(4)
c1.metric("Gasto total", money(df_f["monto"].sum()))
c2.metric("Promedio", money(df_f["monto"].mean()))
c3.metric("Máximo", money(df_f["monto"].max()))
c4.metric("Registros", int(len(df_f)))

st.divider()

nivel = st.radio(
    "Nivel de análisis",
    ["Mensual", "Semanal"],
    horizontal=True,
    index=0
)

df_p = make_period_labels(df_f, nivel=nivel)

resumen_periodo = (
    df_p.groupby(["periodo", "periodo_etq"], as_index=False)
        .agg(
            gasto_total=("monto", "sum"),
            gasto_promedio=("monto", "mean"),
            movimientos=("monto", "count")
        )
        .sort_values("periodo")
)

resumen_categoria = (
    df_f.groupby("categoria")["monto"].sum().sort_values(ascending=False)
)

unidad = "mes" if nivel == "Mensual" else "semana"
insights = build_insights(resumen_periodo, resumen_categoria, unidad=unidad)

colA, colB = st.columns(2)

with colA:
    st.markdown(f"## Evolución del gasto {nivel.lower()}")

    fig, ax = plt.subplots(figsize=(7.4, 4.2))
    ax.plot(resumen_periodo["periodo_etq"], resumen_periodo["gasto_total"], marker="o", linewidth=2)
    ax.set_xlabel("Periodo")
    ax.set_ylabel("Monto ($)")
    ax.set_title(f"Gasto total por {nivel.lower()}")
    plt.xticks(rotation=35, ha="right")

    style_dark_matplotlib(ax)
    plt.tight_layout()
    st.pyplot(fig, use_container_width=True)

with colB:
    st.markdown("## Gasto total por categoría")

    fig2, ax2 = plt.subplots(figsize=(7.4, 4.2))
    ax2.bar(resumen_categoria.index.astype(str), resumen_categoria.values.astype(float))
    ax2.set_xlabel("Categoría")
    ax2.set_ylabel("Monto ($)")
    ax2.set_title("Ranking de gasto por categoría")
    plt.xticks(rotation=35, ha="right")

    style_dark_matplotlib(ax2)
    plt.tight_layout()
    st.pyplot(fig2, use_container_width=True)

st.markdown(
    f"""
    <div class="insights-grid">
      <div class="insight-card">
        <div class="insight-title">Interpretación automática — {nivel}</div>
        <p class="insight-text">{insights["periodo"]}</p>
      </div>
      <div class="insight-card">
        <div class="insight-title">Interpretación automática — Categoría</div>
        <p class="insight-text">{insights["categoria"]}</p>
      </div>
    </div>
    """,
    unsafe_allow_html=True
)

st.divider()

st.markdown("## Resumen del periodo (tabla)")

resumen_show = pd.DataFrame({
    "periodo": resumen_periodo["periodo_etq"].astype(str),
    "gasto_total": resumen_periodo["gasto_total"],
    "gasto_promedio": resumen_periodo["gasto_promedio"],
    "movimientos": resumen_periodo["movimientos"],
})

st.dataframe(
    format_table_for_display(resumen_show, money_cols=["gasto_total", "gasto_promedio"]),
    use_container_width=True,
    height=260
)

st.markdown("## Detalle de movimientos filtrados")
detalle = df_f.copy()
detalle["fecha"] = detalle["fecha"].dt.strftime("%Y-%m-%d")

st.dataframe(
    format_table_for_display(detalle, money_cols=["monto"]),
    use_container_width=True,
    height=420
)

#st.download_button(
#    "⬇ Descargar CSV (filtrado)",
#    data=df_f.to_csv(index=False).encode("utf-8"),
#    file_name="reporte_descriptivo_filtrado.csv",
#    mime="text/csv"
#)

st.caption("Generar Word con detalle completo puede tardar si hay muchos registros.")

if "word_bytes_descriptiva" not in st.session_state:
    st.session_state.word_bytes_descriptiva = None

if st.button("📄 Generar Word (detalle completo)"):
    word_df = df_f.copy()
    word_df["fecha"] = word_df["fecha"].dt.strftime("%Y-%m-%d")

    st.session_state.word_bytes_descriptiva = build_word_report_full(
        df_filtrado=word_df,
        resumen_periodo_show=resumen_show,
        resumen_categoria=resumen_categoria,
        insights=insights,
        nivel=nivel,
        start_date=start,
        end_date=end,
        fig_evolucion=fig,
        fig_categorias=fig2
    )

if st.session_state.word_bytes_descriptiva:
    st.download_button(
        "⬇ Descargar Reporte Word (.docx)",
        data=st.session_state.word_bytes_descriptiva,
        file_name="reporte_descriptivo_completo.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.caption("© 2025 Portal de Analítica | Módulo Descriptivo")