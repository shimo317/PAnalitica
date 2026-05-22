import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt

from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import silhouette_score

from pathlib import Path
from io import BytesIO
from datetime import datetime

# ✅ Toggle global (botón + abrir/cerrar sidebar)
from ui.sidebar_toggle import sidebar_toggle


# ==========================================
# CONFIG
# ==========================================
st.set_page_config(
    page_title="Minería / Segmentación",
    layout="wide",
    initial_sidebar_state="expanded"
)

from auth import init_db, load_authenticated_user, logout_user

init_db()
user = load_authenticated_user()

if not user:
    st.warning("Tu sesión ha expirado o no has iniciado sesión.")
    st.switch_page("Home.py")

if user["role"] not in ["admin", "analista", "usuario"]:
    st.error("No tienes permisos de rol para acceder a esta sección.")
    st.stop()

if "Minería / Segmentación" not in user.get("allowed_modules", []):
    st.error("🔒 Módulo bloqueado. No tienes autorización para acceder a 'Minería / Segmentación'.")
    st.stop()

# Bootstrap icons
st.markdown(
    '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.1/font/bootstrap-icons.css">',
    unsafe_allow_html=True
)

# ==========================================
# Cargar CSS (global + módulo)
# ==========================================
base_dir = Path(__file__).resolve().parents[1]  # .../portal_analitica
global_css = base_dir / "styles" / "global.css"
module_css = base_dir / "styles" / "mineria_segmentacion.css"

css_text = ""
if global_css.exists():
    css_text += global_css.read_text(encoding="utf-8") + "\n"
if module_css.exists():
    css_text += module_css.read_text(encoding="utf-8") + "\n"

if css_text.strip():
    st.markdown(f"<style>{css_text}</style>", unsafe_allow_html=True)
else:
    st.warning("No se encontró CSS (global.css / mineria_segmentacion.css).")

# ✅ Toggle después del CSS
sidebar_toggle(user)

# ==========================================
# Helpers
# ==========================================
def style_dark_matplotlib(ax):
    ax.set_facecolor((0, 0, 0, 0))
    ax.figure.set_facecolor((0, 0, 0, 0))
    ax.tick_params(colors="white")
    ax.xaxis.label.set_color("white")
    ax.yaxis.label.set_color("white")
    ax.title.set_color("white")
    ax.grid(True, alpha=0.18)
    for spine in ax.spines.values():
        spine.set_alpha(0.25)
        spine.set_color("white")

def clamp(v, lo, hi):
    return max(lo, min(hi, v))

def build_interpretation(k, sil, n_used, cols_sel, warnings):
    """
    Regresa: nivel, exec_line, bullets(list), next_step, note
    """
    bullets = []
    warn_count = len(warnings)

    if sil is None:
        nivel = "Regular"
        bullets.append("<b>Silhouette:</b> no se pudo calcular (pocos datos o un solo cluster).")
    else:
        if sil >= 0.50:
            nivel = "Buena"
            bullets.append(f"✅ <b>Separación de clusters buena</b> (Silhouette {sil:.2f}).")
        elif sil >= 0.25:
            nivel = "Regular"
            bullets.append(f"⚠️ <b>Separación aceptable</b> (Silhouette {sil:.2f}).")
        else:
            nivel = "Riesgosa"
            bullets.append(f"❗ <b>Separación baja</b> (Silhouette {sil:.2f}). Considera cambiar variables o k.")

    bullets.append(f"<b>k seleccionado:</b> {k} segmentos.")
    bullets.append(f"<b>Registros usados:</b> {n_used} (después de limpiar nulos).")
    bullets.append(f"<b>Variables:</b> {', '.join(cols_sel)}.")

    if warn_count > 0:
        bullets.append(f"⚠️ <b>Advertencias:</b> {warn_count} (revisa para mejorar estabilidad).")
        for w in warnings[:3]:
            bullets.append(f"• {w}")
        if warn_count > 3:
            bullets.append("• (Hay más advertencias; revisa el panel de warnings).")
    else:
        bullets.append("✅ <b>Sin advertencias críticas:</b> configuración estable.")

    if nivel == "Buena":
        next_step = "Ya puedes pasar a <b>Interpretación de segmentos</b> y acciones: campañas, perfiles, reglas de negocio."
    elif nivel == "Regular":
        next_step = "Prueba <b>otro k</b> (usa Elbow + Silhouette), o añade/quita variables numéricas."
    else:
        next_step = "Antes de usar estos segmentos, ajusta variables/k; si no, los grupos serán poco confiables."

    exec_line = (
        f"Resumen: k={k} | Registros={n_used} | "
        f"Silhouette={'N/A' if sil is None else f'{sil:.2f}'} | "
        f"Variables={len(cols_sel)}."
    )

    note = (
        "K-Means es sensible a escala y outliers: por eso aquí se aplica StandardScaler "
        "y se recomienda validar k con Elbow/Silhouette."
    )
    return nivel, exec_line, bullets, next_step, note


def fig_to_png_bytes(fig) -> BytesIO:
    bio = BytesIO()
    fig.savefig(bio, format="png", dpi=200, bbox_inches="tight")
    bio.seek(0)
    return bio


def df_to_docx_table(doc, df: pd.DataFrame, title: str, note: str = ""):
    """
    Inserta tabla COMPLETA (sin recorte) en Word.
    Ojo: si hay miles de filas, puede tardar (pero tú pediste TODO).
    """
    from docx.shared import Inches

    doc.add_heading(title, level=2)
    if note:
        p = doc.add_paragraph(note)
        p.runs[0].italic = True

    # Convertir a strings seguros
    safe = df.copy()
    safe.columns = [str(c) for c in safe.columns]
    for c in safe.columns:
        safe[c] = safe[c].astype(str)

    rows, cols = safe.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = "Table Grid"

    # header
    for j, col in enumerate(safe.columns):
        table.cell(0, j).text = str(col)

    # body
    for i in range(rows):
        for j in range(cols):
            table.cell(i + 1, j).text = safe.iat[i, j]


def build_word_report_minera(
    df_original: pd.DataFrame,
    cols_sel,
    k,
    n_used,
    sil,
    seg_df,
    centroids_df,
    summary_df,
    warnings,
    nivel,
    exec_line,
    bullets,
    next_step,
    note,
    figs_dict: dict,
):
    """
    Genera Word (.docx) con TODO: info + tablas completas + imágenes.
    """
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception:
        return None, "Falta python-docx. Instala: pip install python-docx"

    doc = Document()

    # Título
    doc.add_heading("Reporte — Minería de Datos: Segmentación (K-Means)", level=1)
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Resumen ejecutivo
    doc.add_heading("Resumen ejecutivo", level=2)
    doc.add_paragraph(exec_line)
    doc.add_paragraph(f"Nivel: {nivel}")

    doc.add_heading("Interpretación automática", level=2)
    # bullets (limpios de HTML básico)
    def strip_html(s: str) -> str:
        return (s.replace("<b>", "").replace("</b>", "")
                .replace("✅", "").replace("⚠️", "").replace("❗", "").strip())

    for b in bullets:
        b2 = strip_html(b)
        if b2.startswith("•"):
            doc.add_paragraph(b2, style="List Bullet")
        else:
            doc.add_paragraph(b2, style="List Bullet")

    doc.add_paragraph(f"Siguiente paso: {strip_html(next_step)}")
    doc.add_paragraph(f"Nota: {strip_html(note)}")

    # Advertencias
    doc.add_heading("Advertencias", level=2)
    if warnings:
        for w in warnings:
            doc.add_paragraph(w, style="List Bullet")
    else:
        doc.add_paragraph("Sin advertencias críticas.")

    # Parámetros y variables
    doc.add_heading("Configuración", level=2)
    doc.add_paragraph(f"Variables seleccionadas: {', '.join([str(c) for c in cols_sel])}")
    doc.add_paragraph(f"k (segmentos): {int(k)}")
    doc.add_paragraph(f"Registros usados: {int(n_used)}")
    doc.add_paragraph(f"Silhouette: {'N/A' if sil is None else f'{sil:.3f}'}")

    # Tablas
    df_to_docx_table(doc, seg_df, "Tamaño de segmentos")
    df_to_docx_table(doc, centroids_df, "Centroides (perfil promedio por segmento)")
    df_to_docx_table(doc, summary_df, "Promedios por segmento (interpretación de perfiles)")

    # Dataset completo segmentado (TODO)
    doc.add_heading("Detalle completo (datos segmentados)", level=2)
    doc.add_paragraph(
        "Contiene todas las filas usadas para clustering (ya sin nulos en variables seleccionadas) "
        "con su segmento asignado."
    )
    df_to_docx_table(doc, df_original, "Tabla: datos usados + segmento")

    # Imágenes
    doc.add_heading("Gráficas", level=2)
    for name, bio in figs_dict.items():
        doc.add_paragraph(name)
        try:
            doc.add_picture(bio, width=Inches(6.5))
        except Exception:
            # fallback: si alguna versión no acepta BytesIO, no revienta
            doc.add_paragraph(f"(No se pudo insertar la imagen: {name})")

    # Guardar a bytes
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue(), None


# ==========================================
# HERO
# ==========================================
st.markdown(
    """
    <div class="hero-wrap">
        <div class="hero-title">Minería de Datos: <span class="accent">Segmentación</span></div>
        <div class="hero-sub">
            Descubre patrones y segmenta registros mediante <b>clustering (K-Means)</b>.
            Ideal para encontrar <b>grupos de comportamiento</b> en datos numéricos.
        </div>
    </div>
    <div class="sep"></div>
    """,
    unsafe_allow_html=True
)

# ==========================================
# SIDEBAR: guía + configuración
# ==========================================
st.sidebar.markdown(
    """
    <div class="sidebar-help">
        <h4>¿Cómo usar este módulo?</h4>
        <ul>
            <li>Sube tu archivo Excel o CSV.</li>
            <li>Selecciona <b>mínimo 2 variables numéricas</b>.</li>
            <li>Define <b>k</b> (segmentos) y valida con <b>Elbow</b> / <b>Silhouette</b>.</li>
            <li>Interpreta centroides: son el “perfil promedio” de cada grupo.</li>
            <li>Descarga el reporte (CSV / Word / TXT).</li>
        </ul>
        <div class="sidebar-tip">Tip: Segmentación es más confiable cuando escalas variables y evitas outliers extremos.</div>
    </div>
    <div class="sidebar-sep"></div>
    <h4 class="sidebar-title">Configuración</h4>
    """,
    unsafe_allow_html=True
)

preview_rows = st.sidebar.number_input(
    "Filas a previsualizar",
    min_value=5,
    max_value=200,
    value=30,
    step=5
)

advanced = st.sidebar.toggle("Modo avanzado", value=False)

# Parámetros base (seguros)
k_min, k_max = 2, 10
k_default = 3
random_state = 42
n_init = 10

show_elbow = True
show_silhouette_panel = True
elbow_max_k = 10

if advanced:
    st.sidebar.markdown("### Parámetros K-Means")
    k_default = st.sidebar.slider("k por defecto", 2, 10, int(k_default), 1)
    k_max = st.sidebar.slider("k máximo", 3, 20, int(k_max), 1)
    n_init = st.sidebar.slider("n_init", 5, 30, int(n_init), 1)
    random_state = st.sidebar.number_input("random_state", value=int(random_state), step=1)

    st.sidebar.markdown("### Validación de k")
    show_elbow = st.sidebar.toggle("Mostrar Elbow (inercia)", value=True)
    show_silhouette_panel = st.sidebar.toggle("Mostrar Silhouette", value=True)
    elbow_max_k = st.sidebar.slider("Elbow: k máximo", 3, 20, int(min(elbow_max_k, k_max)), 1)

# ==========================================
# PANEL REQUISITOS
# ==========================================
st.markdown(
    """
    <div class="panel">
        <div style="display:flex; gap:12px;">
            <div class="panel-icon"><i class="bi bi-cpu"></i></div>
            <div>
                <div class="panel-title">Requisitos del archivo</div>
                <small>
                    <b>Formato:</b> Excel (.xlsx) o CSV (.csv) &nbsp;&nbsp;|&nbsp;&nbsp;
                    <b>Regla:</b> al menos <b>2 columnas numéricas</b> para segmentar
                </small>
            </div>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

# ==========================================
# CARGA CSV
# ==========================================
#file = st.file_uploader(
#    "Sube tu archivo Excel (debe tener al menos 2 columnas numéricas)",
#    type=["xlsx"]
#)

file = st.file_uploader(
    "Sube tu archivo (Excel o CSV)", 
    type=["xlsx", "csv"]
)

if not file:
    st.info("Sube un archivo Excel o CSV para comenzar.")
    st.stop()

#df = None
#read_errors = []

#for kwargs in [
#    dict(sep=",", encoding="utf-8"),
#    dict(sep=";", encoding="utf-8"),
#    dict(sep=",", encoding="latin-1"),
#    dict(sep=";", encoding="latin-1"),
#]:
#    try:
#        file.seek(0)
#        df = pd.read_csv(file, **kwargs)
#        break
#    except Exception as e:
#        read_errors.append(str(e))

# if df is None:
#    st.error("No se pudo leer el CSV con separadores comunes (, ;) y encodings (utf-8/latin-1).")
#    st.code("\n\n".join(read_errors[:3]))
#    st.stop()

#try:
#    file.seek(0)
#    df = pd.read_excel(file)
#except Exception as e:
#    st.error(f"No se pudo leer el archivo Excel. Error: {e}")
#    st.stop()

df = None
error_msg = None

try:
    file.seek(0)

    if file.name.endswith(".xlsx"):
        df = pd.read_excel(file)

    elif file.name.endswith(".csv"):
        # Intentos comunes (manejo realista de CSV)
        for kwargs in [
            dict(sep=",", encoding="utf-8"),
            dict(sep=";", encoding="utf-8"),
            dict(sep=",", encoding="latin-1"),
            dict(sep=";", encoding="latin-1"),
        ]:
            try:
                file.seek(0)
                df = pd.read_csv(file, **kwargs)
                break
            except Exception:
                continue

        if df is None:
            raise ValueError("No se pudo leer el CSV con separadores comunes.")

    else:
        raise ValueError("Formato no soportado")

except Exception as e:
    st.error(f"No se pudo leer el archivo. Error: {e}")
    st.stop()

if df.empty:
    st.warning("El archivo está vacío o no contiene registros.")
    st.stop()

df.columns = [str(c).strip() for c in df.columns]

# ==========================================
# VISTA PREVIA
# ==========================================
st.markdown('<div class="section-title">Vista previa del dataset</div>', unsafe_allow_html=True)
st.dataframe(df.head(int(preview_rows)), use_container_width=True)
st.divider()

# ==========================================
# Selección de columnas numéricas
# ==========================================
num_cols = df.select_dtypes(include=[np.number]).columns.tolist()

if len(num_cols) < 2:
    st.error("Se necesitan al menos **2 columnas numéricas** para hacer clustering.")
    st.write("Columnas numéricas detectadas:", num_cols)
    st.stop()

cols_sel = st.sidebar.multiselect(
    "Variables numéricas (mínimo 2)",
    options=num_cols,
    default=num_cols[: min(3, len(num_cols))]
)

if len(cols_sel) < 2:
    st.warning("Selecciona mínimo 2 columnas numéricas.")
    st.stop()

st.markdown('<div class="section-title">1) Variables seleccionadas</div>', unsafe_allow_html=True)
st.write(f"Columnas: **{', '.join(cols_sel)}**")

# Datos para clustering (limpieza simple)
data = df[cols_sel].copy()
n_total = len(data)
n_null_rows = int(data.isna().any(axis=1).sum())

data = data.dropna()
n_used = len(data)

warnings = []
if n_used < 10:
    warnings.append("Muy pocos registros después de eliminar nulos. El clustering puede ser inestable.")
if n_total > 0 and n_used < (0.60 * n_total):
    warnings.append("Se eliminaron muchas filas por nulos. Considera limpiar/imputar en Calidad de Datos.")

st.write(
    f"Registros totales: **{n_total}** | "
    f"Filas con nulos en variables: **{n_null_rows}** | "
    f"Registros usados: **{n_used}**"
)

if n_used < 5:
    st.warning("Muy pocos registros válidos después de eliminar nulos. Sube un dataset con más filas.")
    st.stop()

# Parámetro K
k_default = int(clamp(k_default, k_min, k_max))
k = st.sidebar.slider(
    "Número de segmentos (k)",
    min_value=int(k_min),
    max_value=int(k_max),
    value=int(k_default),
    step=1
)

# ==========================================
# 2) Validación de k (Elbow + Silhouette)
# ==========================================
st.markdown('<div class="section-title">2) Validación del número de segmentos</div>', unsafe_allow_html=True)

scaler = StandardScaler()
X = scaler.fit_transform(data.values)

fig_elbow = None
fig_sil = None

# Elbow (inercia)
if show_elbow:
    max_k = int(clamp(elbow_max_k, 3, k_max))
    ks = list(range(2, max_k + 1))
    inertias = []

    for kk in ks:
        km = KMeans(n_clusters=int(kk), random_state=int(random_state), n_init=int(n_init))
        km.fit(X)
        inertias.append(float(km.inertia_))

    fig_elbow, ax_elbow = plt.subplots(figsize=(7.6, 4.0))
    ax_elbow.plot(ks, inertias, marker="o")
    ax_elbow.set_title("Elbow (Inercia vs k)")
    ax_elbow.set_xlabel("k (número de segmentos)")
    ax_elbow.set_ylabel("Inercia (SSE)")
    style_dark_matplotlib(ax_elbow)
    st.pyplot(fig_elbow, use_container_width=True)
    st.caption("Tip: busca el “codo” donde deja de mejorar mucho al aumentar k.")

# Silhouette (panel)
sil_k_best = None
if show_silhouette_panel:
    if n_used >= 10:
        max_sil_k = int(min(k_max, 12))
        ks2 = list(range(2, max_sil_k + 1))
        sils = []

        for kk in ks2:
            try:
                km = KMeans(n_clusters=int(kk), random_state=int(random_state), n_init=int(n_init))
                labels_tmp = km.fit_predict(X)
                s = silhouette_score(X, labels_tmp)
                sils.append(float(s))
            except Exception:
                sils.append(np.nan)

        fig_sil, ax_sil = plt.subplots(figsize=(7.6, 4.0))
        ax_sil.plot(ks2, sils, marker="o")
        ax_sil.set_title("Silhouette Score vs k")
        ax_sil.set_xlabel("k (número de segmentos)")
        ax_sil.set_ylabel("Silhouette")
        style_dark_matplotlib(ax_sil)
        st.pyplot(fig_sil, use_container_width=True)

        if np.isfinite(np.nanmax(sils)):
            best_idx = int(np.nanargmax(sils))
            sil_k_best = int(ks2[best_idx])
            st.caption(f"Mejor k (Silhouette): **{sil_k_best}** (orientativo).")
    else:
        st.info("Silhouette requiere más datos (recomendado >= 10).")

st.divider()

# ==========================================
# 3) Entrenamiento K-Means (k seleccionado)
# ==========================================
st.markdown('<div class="section-title">3) Clustering con K-Means</div>', unsafe_allow_html=True)

model = KMeans(n_clusters=int(k), random_state=int(random_state), n_init=int(n_init))
labels = model.fit_predict(X)

# silhouette del k elegido
sil = None
try:
    if len(np.unique(labels)) > 1 and n_used >= 10:
        sil = float(silhouette_score(X, labels))
except Exception:
    sil = None

data_out = data.copy()
data_out["segmento"] = labels

c1, c2, c3 = st.columns(3)
c1.metric("k (segmentos)", int(k))
c2.metric("Registros usados", int(n_used))
c3.metric("Silhouette", "N/A" if sil is None else f"{sil:.2f}")

# Tamaño segmentos
counts = data_out["segmento"].value_counts().sort_index()
seg_df = pd.DataFrame({"segmento": counts.index, "registros": counts.values}).sort_values("segmento")
st.write("**Tamaño de segmentos:**")
st.dataframe(seg_df, use_container_width=True)

# Centroides en unidades originales
centroids_scaled = model.cluster_centers_
centroids = scaler.inverse_transform(centroids_scaled)
centroids_df = pd.DataFrame(centroids, columns=cols_sel)
centroids_df["segmento"] = range(int(k))
centroids_df = centroids_df[["segmento"] + cols_sel]

st.write("**Centroides (perfil promedio por segmento):**")
st.dataframe(centroids_df, use_container_width=True)

# Resumen por segmento (promedios)
summary_df = data_out.groupby("segmento")[cols_sel].mean(numeric_only=True).reset_index()
st.write("**Promedios por segmento (para interpretar perfiles):**")
st.dataframe(summary_df, use_container_width=True)

st.divider()

# ==========================================
# 4) Visualización 2D + centroides
# ==========================================
st.markdown('<div class="section-title">4) Visualización (2D)</div>', unsafe_allow_html=True)
st.write("Elige 2 variables para graficar y observar cómo se separan los segmentos:")

x_col = st.selectbox("Eje X", options=cols_sel, index=0)
y_col = st.selectbox("Eje Y", options=cols_sel, index=1 if len(cols_sel) > 1 else 0)

cent_2d = centroids_df.set_index("segmento")[[x_col, y_col]].copy()

fig2d, ax2d = plt.subplots(figsize=(8.2, 4.6))
ax2d.scatter(
    data_out[x_col],
    data_out[y_col],
    c=data_out["segmento"],
    alpha=0.9
)
ax2d.scatter(
    cent_2d[x_col],
    cent_2d[y_col],
    marker="X",
    s=180,
    edgecolors="white",
    linewidths=1.2
)

ax2d.set_xlabel(x_col)
ax2d.set_ylabel(y_col)
ax2d.set_title("Segmentación (colores) + centroides (X)")
style_dark_matplotlib(ax2d)
plt.tight_layout()
st.pyplot(fig2d, use_container_width=True)

st.caption(
    "Interpretación: cada color representa un segmento. "
    "Los centroides (X) son el perfil promedio de cada grupo."
)

# ==========================================
# 5) Interpretación automática (PRO CARD)
# ==========================================
st.markdown('<div class="section-title">5) Interpretación automática</div>', unsafe_allow_html=True)

nivel, exec_line, bullets, next_step, note = build_interpretation(
    k=int(k),
    sil=sil,
    n_used=int(n_used),
    cols_sel=cols_sel,
    warnings=warnings
)

badge_cls = "badge-ok" if nivel == "Buena" else ("badge-warn" if nivel == "Regular" else "badge-bad")
bullets_html = "".join([f"<li>{b}</li>" for b in bullets])

card_html = f"""
<div class="interpret-card">
  <div class="interpret-head">
    <div class="interpret-title">Interpretación automática</div>
    <span class="badge {badge_cls}">{nivel}</span>
  </div>

  <div class="interpret-exec">{exec_line}</div>

  <div class="interpret-text">
    Este módulo realiza <b>segmentación</b> agrupando registros parecidos en <b>k</b> grupos.
    El objetivo es identificar <b>patrones</b> y perfiles (por ejemplo: gasto alto/medio/bajo, comportamiento, etc.).
  </div>

  <ul class="interpret-bullets">
    {bullets_html}
  </ul>

  <div class="interpret-tip">
    <span class="tip-label">Siguiente paso:</span> {next_step}
  </div>

  <div class="interpret-note">{note}</div>
</div>
"""
st.markdown(card_html, unsafe_allow_html=True)

# Warnings claros
if warnings:
    st.markdown('<div class="section-title">Advertencias</div>', unsafe_allow_html=True)
    for w in warnings:
        st.warning(w)

# ==========================================
# DESCARGAS (PRO)
# ==========================================
st.divider()
st.markdown('<div class="section-title">Descargar reporte</div>', unsafe_allow_html=True)

st.markdown(
    """
    <div class="download-card">
      <div class="download-head">
        <div class="download-icon"><i class="bi bi-download"></i></div>
        <div>
          <div class="download-title">Exportaciones</div>
          <div class="download-sub">Incluye CSV, centroides, resumen TXT y reporte Word con TODO.</div>
        </div>
      </div>
    </div>
    """,
    unsafe_allow_html=True
)

# 1) CSV con segmentos
csv_out = data_out.reset_index(drop=True).to_csv(index=False).encode("utf-8")

# 2) CSV centroides
csv_cent = centroids_df.to_csv(index=False).encode("utf-8")

# 3) TXT resumen
txt = (
    f"{exec_line}\n"
    f"Nivel: {nivel}\n"
    f"Variables: {', '.join(cols_sel)}\n"
    f"Silhouette: {'N/A' if sil is None else f'{sil:.3f}'}\n"
    f"Segmentos (k): {k}\n"
    f"Registros usados: {n_used}\n\n"
    "Interpretación:\n- " + "\n- ".join([b.replace('<b>','').replace('</b>','') for b in bullets]) + "\n\n"
    f"Siguiente paso: {next_step.replace('<b>','').replace('</b>','')}\n"
    f"Nota: {note.replace('<b>','').replace('</b>','')}\n"
)
txt_bytes = txt.encode("utf-8")

# 4) WORD con TODO (tablas + imágenes)
st.caption("Generar Word con detalle completo puede tardar si hay muchos registros.")

figs_dict = {}
if fig_elbow is not None:
    figs_dict["Elbow (Inercia vs k)"] = fig_to_png_bytes(fig_elbow)
if fig_sil is not None:
    figs_dict["Silhouette Score vs k"] = fig_to_png_bytes(fig_sil)
figs_dict["Segmentación 2D + centroides"] = fig_to_png_bytes(fig2d)

# Para el Word: usamos SOLO las filas usadas (data_out) porque son las que realmente se segmentaron
word_bytes, word_err = build_word_report_minera(
    df_original=data_out.reset_index(drop=True),
    cols_sel=cols_sel,
    k=k,
    n_used=n_used,
    sil=sil,
    seg_df=seg_df,
    centroids_df=centroids_df,
    summary_df=summary_df,
    warnings=warnings,
    nivel=nivel,
    exec_line=exec_line,
    bullets=bullets,
    next_step=next_step,
    note=note,
    figs_dict=figs_dict
)

colD1, colD2 = st.columns(2)

with colD1:
    st.download_button(
        "⬇ Descargar CSV (datos con segmento)",
        data=csv_out,
        file_name="reporte_mineria_segmentacion.csv",
        mime="text/csv",
        key="dl_seg"
    )
    st.download_button(
        "⬇ Descargar CSV (centroides)",
        data=csv_cent,
        file_name="reporte_mineria_centroides.csv",
        mime="text/csv",
        key="dl_cent"
    )

with colD2:
    st.download_button(
        "⬇ Descargar TXT (resumen)",
        data=txt_bytes,
        file_name="reporte_mineria_resumen.txt",
        mime="text/plain",
        key="dl_txt"
    )

    if word_err:
        st.warning(word_err)
    else:
        st.download_button(
            "⬇ Descargar REPORTE WORD (.docx)",
            data=word_bytes,
            file_name="reporte_mineria_segmentacion.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_word"
        )

st.caption("© 2025 Portal de Analítica | Módulo Minería / Segmentación")
