import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import re
import unicodedata
from pathlib import Path
from io import BytesIO
from datetime import datetime

# ✅ Word
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ✅ Toggle global (botón + abrir/cerrar sidebar)
from ui.sidebar_toggle import sidebar_toggle


# ==========================================
# CONFIG
# ==========================================
st.set_page_config(
    page_title="Calidad de Datos",
    layout="wide",
    initial_sidebar_state="expanded"
)

from auth import init_db, load_authenticated_user, logout_user

init_db()
user = load_authenticated_user()

if not user:
    st.warning("Tu sesión ha expirado o no has iniciado sesión.")
    st.switch_page("Home.py")

if user["role"] not in ["admin", "analista", "usuario"]:
    st.error("No tienes permisos de rol para acceder a esta sección.")
    st.stop()

if "Calidad de Datos" not in user.get("allowed_modules", []):
    st.error("🔒 Módulo bloqueado. No tienes autorización para acceder a 'Calidad de Datos'.")
    st.stop()

# Bootstrap icons
st.markdown(
    '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.1/font/bootstrap-icons.css">',
    unsafe_allow_html=True
)

# ==========================================
# Cargar CSS (global + módulo)
# ==========================================
base_dir = Path(__file__).resolve().parents[1]  # .../portal_analitica
global_css = base_dir / "styles" / "global.css"
module_css = base_dir / "styles" / "calidad_datos.css"

css_text = ""
if global_css.exists():
    css_text += global_css.read_text(encoding="utf-8") + "\n"
if module_css.exists():
    css_text += module_css.read_text(encoding="utf-8") + "\n"

if css_text.strip():
    st.markdown(f"<style>{css_text}</style>", unsafe_allow_html=True)
else:
    st.warning("No se encontró CSS (global.css / calidad_datos.css).")

# ✅ Toggle después del CSS
sidebar_toggle(user)

# ==========================================
# Helpers
# ==========================================
def style_dark_matplotlib(ax):
    ax.set_facecolor((0, 0, 0, 0))
    ax.figure.set_facecolor((0, 0, 0, 0))
    ax.tick_params(colors="white")
    ax.xaxis.label.set_color("white")
    ax.yaxis.label.set_color("white")
    ax.title.set_color("white")
    ax.grid(True, alpha=0.18)
    for spine in ax.spines.values():
        spine.set_alpha(0.25)
        spine.set_color("white")

def score_level(score: float):
    if score >= 90:
        return "Excelente"
    if score >= 75:
        return "Buena"
    if score >= 60:
        return "Regular"
    return "Riesgosa"

def clamp(v, lo, hi):
    return max(lo, min(hi, v))

def safe_pct(n, total):
    return 0.0 if total <= 0 else (n / total) * 100

def compute_outliers_iqr(series: pd.Series, factor: float):
    s = pd.to_numeric(series, errors="coerce").dropna()
    if len(s) < 5:
        return 0, 0.0
    q1 = float(s.quantile(0.25))
    q3 = float(s.quantile(0.75))
    iqr = q3 - q1
    if iqr == 0:
        return 0, 0.0
    lower = q1 - factor * iqr
    upper = q3 + factor * iqr
    out_count = int(((s < lower) | (s > upper)).sum())
    out_pct = safe_pct(out_count, len(series))
    return out_count, round(out_pct, 2)

def build_quality_explain(score, nivel, dup_count, dup_pct, avg_null_pct, avg_out_pct):
    bullets = []

    if avg_null_pct > 10:
        bullets.append("<b>Nulos altos:</b> revisa columnas clave y define estrategia (imputar/eliminar/validar).")
    elif avg_null_pct > 0:
        bullets.append("<b>Nulos detectados:</b> están en nivel manejable; revisa si afectan columnas clave.")
    else:
        bullets.append("✅ <b>Sin nulos relevantes:</b> buena señal para modelado.")

    if dup_count > 0:
        bullets.append(f"<b>Duplicados:</b> hay {dup_count} filas duplicadas ({dup_pct:.2f}%). Elimínalos para evitar sesgos.")
    else:
        bullets.append("✅ <b>Sin duplicados:</b> evita conteos inflados y sesgos.")

    if avg_out_pct > 5:
        bullets.append("<b>Outliers:</b> hay valores extremos; valida si son errores o casos reales (pueden mover modelos).")
    elif avg_out_pct > 0:
        bullets.append("<b>Outliers ligeros:</b> revisa si son esperables en tu contexto.")
    else:
        bullets.append("✅ <b>Sin outliers relevantes:</b> datos más estables para modelos.")

    if nivel in ["Excelente", "Buena"]:
        next_step = "Tu dataset está bastante listo. Limpia lo mínimo (si aplica) y pasa a Predictiva / Minería."
    elif nivel == "Regular":
        next_step = "Haz una limpieza básica (nulos/duplicados/outliers) y vuelve a correr este módulo."
    else:
        next_step = "Primero corrige calidad (nulos/duplicados/tipos/outliers). Si no, Predictiva/Minería dará resultados raros."

    exec_line = (
        f"Resumen: Score {score:.2f}/100 | "
        f"Duplicados: {dup_count} ({dup_pct:.2f}%) | "
        f"Nulos prom.: {avg_null_pct:.2f}% | "
        f"Outliers prom.: {avg_out_pct:.2f}%."
    )
    return exec_line, bullets, next_step


# ==========================
# Word helpers
# ==========================
def fig_to_png_bytes(fig) -> bytes:
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=200, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()

def add_df_to_doc(doc: Document, df: pd.DataFrame, title: str, max_rows: int = 60):
    doc.add_heading(title, level=2)

    if df is None or df.empty:
        p = doc.add_paragraph("Sin datos para mostrar.")
        p.runs[0].italic = True
        return

    show = df.copy()
    if len(show) > max_rows:
        show = show.head(max_rows)
        doc.add_paragraph(f"(Mostrando primeras {max_rows} filas)")

    # Convertir todo a string para evitar errores en docx
    show = show.fillna("").astype(str)

    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(show.columns):
        hdr_cells[j].text = str(col)

    for _, row in show.iterrows():
        row_cells = table.add_row().cells
        for j, val in enumerate(row.values):
            row_cells[j].text = str(val)

def make_quality_docx_bytes(
    df: pd.DataFrame,
    preview_rows: int,
    dtypes_df: pd.DataFrame,
    nulos_df: pd.DataFrame,
    outliers_df: pd.DataFrame,
    dup_count: int,
    dup_pct: float,
    avg_null_pct: float,
    avg_out_pct: float,
    score: float,
    nivel: str,
    exec_line: str,
    bullets_plain: list,
    next_step: str,
    advanced: bool,
    params: dict,
    fig_nulls_png: bytes | None,
    fig_out_png: bytes | None,
):
    doc = Document()

    # Title
    doc.add_heading("Reporte — Calidad de Datos", level=0)
    meta = doc.add_paragraph()
    meta.add_run("Generado: ").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))

    doc.add_paragraph(" ")

    # Resumen ejecutivo
    doc.add_heading("Resumen ejecutivo", level=1)
    p = doc.add_paragraph()
    p.add_run("Score: ").bold = True
    p.add_run(f"{score:.2f}/100")
    p = doc.add_paragraph()
    p.add_run("Nivel: ").bold = True
    p.add_run(nivel)

    doc.add_paragraph(exec_line)

    # Hallazgos (bullets)
    doc.add_heading("Hallazgos principales", level=1)
    for b in bullets_plain:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Siguiente paso recomendado", level=1)
    doc.add_paragraph(next_step)

    # Configuración usada
    doc.add_heading("Configuración usada", level=1)
    doc.add_paragraph(f"Modo avanzado: {'Sí' if advanced else 'No'}")
    for k, v in params.items():
        doc.add_paragraph(f"{k}: {v}")

    # Preview
    add_df_to_doc(doc, df.head(int(preview_rows)), "Vista previa del dataset")

    # Tipos
    add_df_to_doc(doc, dtypes_df, "Tipos de datos detectados")

    # Nulos
    add_df_to_doc(doc, nulos_df, "Valores nulos por columna", max_rows=80)

    # Duplicados
    doc.add_heading("Filas duplicadas", level=2)
    doc.add_paragraph(f"Duplicados detectados: {dup_count} ({dup_pct:.2f}%)")

    # Outliers
    add_df_to_doc(doc, outliers_df, "Outliers (IQR) en columnas numéricas", max_rows=80)

    # Gráficas
    doc.add_heading("Gráficas", level=1)

    if fig_nulls_png:
        doc.add_paragraph("Nulos (Top 12 columnas):")
        doc.add_picture(BytesIO(fig_nulls_png), width=Inches(6.3))

    if fig_out_png:
        doc.add_paragraph("Outliers (Top 12 numéricas):")
        doc.add_picture(BytesIO(fig_out_png), width=Inches(6.3))

    doc.add_paragraph(" ")
    foot = doc.add_paragraph("© 2025 Portal de Analítica | Módulo Calidad de Datos")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ==========================================
# HERO
# ==========================================
st.markdown(
    """
    <div class="hero-wrap">
        <div class="hero-title">Calidad de <span class="accent">Datos</span></div>
        <div class="hero-sub">
            Evalúa tu dataset antes de analítica o minería:
            <b>nulos</b>, <b>duplicados</b>, <b>tipos</b> y <b>outliers</b>.
        </div>
    </div>
    <div class="sep"></div>
    """,
    unsafe_allow_html=True
)

# ==========================================
# SIDEBAR: guía + configuración
# ==========================================
st.sidebar.markdown(
    """
    <div class="sidebar-help">
        <h4>¿Cómo usar este módulo?</h4>
        <ul>
            <li>Sube tu archivo Excel o CSV (cualquier estructura).</li>
            <li>Revisa tipos, nulos, duplicados y outliers.</li>
            <li>Activa <b>Modo avanzado</b> si deseas ajustar sensibilidad.</li>
            <li>Descarga el reporte antes de modelar.</li>
        </ul>
        <div class="sidebar-tip">
            Tip: si Predictiva/Minería da métricas raras, casi siempre la causa está aquí.
        </div>
    </div>
    <div class="sidebar-sep"></div>
    <h4 class="sidebar-title">Configuración</h4>
    """,
    unsafe_allow_html=True
)

preview_rows = st.sidebar.number_input(
    "Filas a previsualizar",
    min_value=5,
    max_value=200,
    value=30,
    step=5
)

advanced = st.sidebar.toggle("Modo avanzado (sensibilidad del score)", value=False)

# Defaults “seguros”
w_null = 1.2
w_dup = 0.7
w_out = 1.9
iqr_factor = 1.7
min_n_numeric = 16

if advanced:
    st.sidebar.markdown("### Sensibilidad (score)")
    w_null = st.sidebar.slider("Peso Nulos", 0.0, 3.0, float(w_null), 0.05)
    w_dup = st.sidebar.slider("Peso Duplicados", 0.0, 3.0, float(w_dup), 0.05)
    w_out = st.sidebar.slider("Peso Outliers", 0.0, 3.0, float(w_out), 0.05)

    st.sidebar.markdown("### Outliers (IQR)")
    iqr_factor = st.sidebar.slider("Factor IQR", 1.0, 3.5, float(iqr_factor), 0.05)
    min_n_numeric = st.sidebar.number_input("Mín. datos por columna numérica", 5, 200, int(min_n_numeric), 1)

# ==========================================
# PANEL REQUISITOS
# ==========================================
st.markdown(
    """
    <div class="panel">
        <div style="display:flex; gap:12px;">
            <div class="panel-icon"><i class="bi bi-shield-check"></i></div>
            <div>
                <div class="panel-title">¿Qué revisa este módulo?</div>
                <small>
                    <b>Entrada:</b> Archivo Excel (.xlsx) o CSV (.csv) con cualquier estructura &nbsp;|&nbsp;
                    <b>Salida:</b> diagnóstico + score + recomendaciones + reporte descargable
                </small>
            </div>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

# ==========================================
# CARGA CSV
# ==========================================
#file = st.file_uploader("Sube tu archivo Excel (cualquier estructura)", type=["xlsx"])

file = st.file_uploader(
    "Sube tu archivo (Excel, CSV o TXT)",
    type=["xlsx", "csv", "txt"]
)

if not file:
    st.info("Sube un archivo para comenzar.")
    st.stop()

try:
    file.seek(0)

    if file.name.endswith(".xlsx"):
        df = pd.read_excel(file)

    elif file.name.endswith(".csv"):
        df = pd.read_csv(file, sep=None, engine="python")

    elif file.name.endswith(".txt"):
        df = pd.read_csv(file, sep=None, engine="python")

    else:
        st.error("Formato no soportado.")
        st.stop()

except Exception as e:
    st.error(f"No se pudo leer el archivo. Error: {e}")
    st.stop()

if df.empty:
    st.warning("El archivo está vacío o no contiene registros.")
    st.stop()

def clean_col(c):
    c = str(c)
    c = unicodedata.normalize('NFKD', c).encode('ascii', 'ignore').decode('utf-8')
    c = c.strip().lower()
    c = re.sub(r'\s+', '_', c)
    return c

df.columns = [clean_col(c) for c in df.columns]

# ==========================================
# VISTA PREVIA
# ==========================================
st.markdown('<div class="section-title">Vista previa del dataset</div>', unsafe_allow_html=True)
st.dataframe(df.head(int(preview_rows)), use_container_width=True)
st.divider()

# ==========================================
# 1) TIPOS DE DATOS
# ==========================================
st.markdown('<div class="section-title">1) Tipos de datos detectados</div>', unsafe_allow_html=True)

dtypes_df = pd.DataFrame({
    "columna": df.columns,
    "tipo_detectado": [str(df[c].dtype) for c in df.columns],
    "valores_unicos": [int(df[c].nunique(dropna=True)) for c in df.columns]
})
st.dataframe(dtypes_df, use_container_width=True)

# ==========================================
# 2) NULOS
# ==========================================
st.markdown('<div class="section-title">2) Valores nulos por columna</div>', unsafe_allow_html=True)
nulos = df.isna().sum()
porc_nulos = (nulos / len(df)) * 100

nulos_df = pd.DataFrame({
    "columna": nulos.index.astype(str),
    "nulos": nulos.values.astype(int),
    "%_nulos": np.round(porc_nulos.values, 2),
}).sort_values("%_nulos", ascending=False)

st.dataframe(nulos_df, use_container_width=True)

# ==========================================
# 3) DUPLICADOS
# ==========================================
st.markdown('<div class="section-title">3) Filas duplicadas</div>', unsafe_allow_html=True)
dup_count = int(df.duplicated().sum())
dup_pct = safe_pct(dup_count, len(df))
st.write(f"**Filas duplicadas detectadas:** {dup_count} ({dup_pct:.2f}%)")

# ==========================================
# 4) OUTLIERS (IQR) EN NUMÉRICAS
# ==========================================
st.markdown('<div class="section-title">4) Outliers (IQR) en columnas numéricas</div>', unsafe_allow_html=True)

num_cols = df.select_dtypes(include=[np.number]).columns.tolist()

out_rows = []
if not num_cols:
    st.info("No se detectaron columnas numéricas para análisis de outliers.")
    outliers_df = pd.DataFrame(columns=["columna", "outliers_detectados", "%_outliers"])
else:
    for col in num_cols:
        series = df[col]
        if pd.to_numeric(series, errors="coerce").dropna().shape[0] < int(min_n_numeric):
            out_rows.append([str(col), 0, 0.0])
            continue
        out_count, out_pct = compute_outliers_iqr(series, float(iqr_factor))
        out_rows.append([str(col), out_count, out_pct])

    outliers_df = (
        pd.DataFrame(out_rows, columns=["columna", "outliers_detectados", "%_outliers"])
        .sort_values("%_outliers", ascending=False)
        .reset_index(drop=True)
    )

st.dataframe(outliers_df, use_container_width=True)
st.divider()

# ==========================================
# 5) SCORE CALIDAD (ESTIMADO)
# ==========================================
st.markdown('<div class="section-title">5) Score de Calidad (estimado)</div>', unsafe_allow_html=True)

avg_null_pct = float(nulos_df["%_nulos"].mean()) if len(nulos_df) else 0.0
avg_out_pct = float(outliers_df["%_outliers"].mean()) if len(outliers_df) else 0.0

penalty = (avg_null_pct * float(w_null)) + (dup_pct * float(w_dup)) + (avg_out_pct * float(w_out))
score = clamp(100.0 - penalty, 0.0, 100.0)
nivel = score_level(score)

c1, c2, c3, c4 = st.columns(4)
c1.metric("Score (0–100)", f"{score:.2f}")
c2.metric("Nivel", nivel)
c3.metric("Nulos promedio", f"{avg_null_pct:.2f}%")
c4.metric("Outliers promedio", f"{avg_out_pct:.2f}%")

# ==========================================
# VISUALIZACIÓN (calidad)
# ==========================================
st.markdown('<div class="section-title">Visualización (calidad)</div>', unsafe_allow_html=True)

g1, g2 = st.columns(2)

fig_nulls_png = None
fig_out_png = None

with g1:
    st.markdown("### % Nulos (Top 12 columnas)")
    top_nulls = nulos_df.head(12).copy()
    fig, ax = plt.subplots(figsize=(7.4, 4.2))
    ax.bar(top_nulls["columna"].astype(str), top_nulls["%_nulos"].astype(float))
    ax.set_ylabel("% nulos")
    ax.set_title("Columnas con más nulos")
    plt.xticks(rotation=35, ha="right")
    style_dark_matplotlib(ax)
    plt.tight_layout()
    st.pyplot(fig, use_container_width=True)
    fig_nulls_png = fig_to_png_bytes(fig)

with g2:
    st.markdown("### % Outliers (Top 12 numéricas)")
    if len(outliers_df) == 0:
        st.info("No hay columnas numéricas.")
    else:
        top_out = outliers_df.head(12).copy()
        fig2, ax2 = plt.subplots(figsize=(7.4, 4.2))
        ax2.bar(top_out["columna"].astype(str), top_out["%_outliers"].astype(float))
        ax2.set_ylabel("% outliers")
        ax2.set_title("Columnas con más outliers")
        plt.xticks(rotation=35, ha="right")
        style_dark_matplotlib(ax2)
        plt.tight_layout()
        st.pyplot(fig2, use_container_width=True)
        fig_out_png = fig_to_png_bytes(fig2)

# ==========================================
# INTERPRETACIÓN PRO (CARD)
# ==========================================
exec_line, bullets, next_step = build_quality_explain(
    score=score,
    nivel=nivel,
    dup_count=dup_count,
    dup_pct=dup_pct,
    avg_null_pct=avg_null_pct,
    avg_out_pct=avg_out_pct
)

badge_class = "badge-ok" if nivel in ["Excelente", "Buena"] else ("badge-warn" if nivel == "Regular" else "badge-bad")
bullets_html = "".join([f"<li>{b}</li>" for b in bullets])

card_html = f"""
<div class="interpret-card">
  <div class="interpret-head">
    <div class="interpret-title">Interpretación automática</div>
    <span class="badge {badge_class}">{nivel}</span>
  </div>

  <div class="interpret-exec">{exec_line}</div>

  <div class="interpret-text">
    Este módulo sirve para detectar problemas típicos de calidad que luego provocan
    métricas raras, modelos inestables o conclusiones equivocadas.
  </div>

  <ul class="interpret-bullets">
    {bullets_html}
  </ul>

  <div class="interpret-tip">
    <span class="tip-label">Siguiente paso:</span> {next_step}
  </div>

  <div class="interpret-note">
    Recomendado <b>antes</b> de Predictiva / Minería para evitar sesgos,
    errores por tipos incorrectos o resultados inestables por datos extremos.
  </div>
</div>
"""
st.markdown(card_html, unsafe_allow_html=True)

# ==========================================
# DESCARGAS
# ==========================================
st.divider()
st.markdown('<div class="section-title">Descargar reporte</div>', unsafe_allow_html=True)

report_cols = pd.DataFrame({
    "columna": dtypes_df["columna"].astype(str),
    "tipo_detectado": dtypes_df["tipo_detectado"].astype(str),
    "valores_unicos": dtypes_df["valores_unicos"].astype(int),
}).merge(
    nulos_df[["columna", "nulos", "%_nulos"]],
    on="columna",
    how="left"
)

if len(outliers_df):
    report_cols = report_cols.merge(outliers_df, on="columna", how="left")
else:
    report_cols["outliers_detectados"] = 0
    report_cols["%_outliers"] = 0.0

report_cols["outliers_detectados"] = report_cols["outliers_detectados"].fillna(0).astype(int)
report_cols["%_outliers"] = report_cols["%_outliers"].fillna(0.0).astype(float)

# Resumen global (constantes)
report_cols["filas_total"] = int(len(df))
report_cols["duplicados_total"] = int(dup_count)
report_cols["duplicados_pct"] = float(round(dup_pct, 2))
report_cols["nulos_promedio_pct"] = float(round(avg_null_pct, 2))
report_cols["outliers_promedio_pct"] = float(round(avg_out_pct, 2))
report_cols["score_calidad"] = float(round(score, 2))
report_cols["nivel"] = nivel

#csv_bytes = report_cols.to_csv(index=False).encode("utf-8")
#st.download_button(
#    "⬇ Descargar CSV (reporte calidad de datos)",
#   data=csv_bytes,
#   file_name="reporte_calidad_datos.csv",
#    mime="text/csv",
#)

# ==========================
# ✅ NUEVO: WORD (TODO EL REPORTE)
# ==========================
params_used = {
    "Filas previsualizadas": int(preview_rows),
    "Peso Nulos (w_null)": float(w_null),
    "Peso Duplicados (w_dup)": float(w_dup),
    "Peso Outliers (w_out)": float(w_out),
    "Factor IQR": float(iqr_factor),
    "Mín. datos por columna numérica": int(min_n_numeric),
}

bullets_plain = [
    b.replace("<b>", "").replace("</b>", "").replace("✅", "•")
    for b in bullets
]

# Card visual (si ya tienes CSS, se verá pro)
st.markdown(
    """
    <div class="word-card">
      <div class="word-title">Reporte Word (completo)</div>
      <div class="word-sub">
        Incluye: vista previa, tipos, nulos, duplicados, outliers, score, interpretación y gráficas.
      </div>
    </div>
    """,
    unsafe_allow_html=True
)

if st.button("📝 Generar Word (reporte completo)"):
    docx_bytes = make_quality_docx_bytes(
        df=df,
        preview_rows=int(preview_rows),
        dtypes_df=dtypes_df,
        nulos_df=nulos_df,
        outliers_df=outliers_df,
        dup_count=int(dup_count),
        dup_pct=float(dup_pct),
        avg_null_pct=float(avg_null_pct),
        avg_out_pct=float(avg_out_pct),
        score=float(score),
        nivel=str(nivel),
        exec_line=str(exec_line),
        bullets_plain=bullets_plain,
        next_step=str(next_step),
        advanced=bool(advanced),
        params=params_used,
        fig_nulls_png=fig_nulls_png,
        fig_out_png=fig_out_png,
    )

    st.download_button(
        "⬇ Descargar Word (Calidad de Datos)",
        data=docx_bytes,
        file_name="reporte_calidad_datos.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

st.caption("© 2025 Portal de Analítica | Módulo Calidad de Datos")
